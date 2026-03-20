"""RAG Chatbot server — F1 Knowledge Base Agent.

Queries MongoDB Atlas vector search for context, then generates
answers via Groq LLM (llama-3.3-70b-versatile).

Usage:
    python pipeline/chat_server.py          # Runs on port 8300
"""

from __future__ import annotations

import logging
import os
import sys
import json
import tempfile
from pathlib import Path

logger = logging.getLogger(__name__)

import numpy as np

from fastapi import BackgroundTasks, FastAPI, HTTPException, Request, UploadFile, File
from fastapi.responses import JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

# Ensure project root is on path
sys.path.insert(0, str(Path(__file__).parent.parent))
# OmniSuite packages
sys.path.insert(0, str(Path(__file__).parent.parent / "omnisuitef1"))

from dotenv import load_dotenv
load_dotenv(Path(__file__).parent.parent / ".env", override=True)

from groq import Groq
from pipeline.vectorstore import get_vector_store
from pipeline.embeddings import NomicEmbedder
from pipeline.model_3d_server import router as model_3d_router, mount_3d_static
from pipeline.omni_health_router import router as omni_health_router
from pipeline.omni_analytics_router import router as omni_analytics_router
from pipeline.omni_rag_router import router as omni_rag_router
from pipeline.omni_kex_router import router as omni_kex_router
from pipeline.omni_doc_router import router as omni_doc_router
from pipeline.omni_data_router import router as omni_data_router
from pipeline.omni_bedding_router import router as omni_bedding_router
from pipeline.omni_vis_router import router as omni_vis_router
from pipeline.omni_dapt_router import router as omni_dapt_router
from pipeline.opponents.server import router as opponents_router, init_profiler_with_db
from pipeline.updater.server import router as updater_router
from pipeline.advantage_router import router as advantage_router
from pipeline.omni_agents_router import router as omni_agents_router
from pipeline.radio_router import router as radio_router
from pipeline.aim_router import router as aim_router

# ── Config ───────────────────────────────────────────────────────────────

GROQ_MODEL = os.getenv("GROQ_REASONING_MODEL", "llama-3.3-70b-versatile")
PORT = int(os.getenv("PORT", os.getenv("API_PORT", "8300")))
USE_OMNIRAG = os.getenv("USE_OMNIRAG", "").lower() in ("1", "true", "yes")

app = FastAPI(title="F1 OmniSense API")
_allowed_origins = os.getenv("ALLOWED_ORIGINS", "http://localhost:5173,http://localhost:3000").split(",")
app.add_middleware(
    CORSMiddleware,
    allow_origins=_allowed_origins,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Rewrite /api/X → /api/local/X (frontend expects short paths, backend uses /api/local/)
from starlette.middleware.base import BaseHTTPMiddleware
from starlette.requests import Request as _Request

class _RewriteMiddleware(BaseHTTPMiddleware):
    _PREFIXES = (
        "/api/openf1/", "/api/jolpica/", "/api/driver_intel/",
        "/api/circuit_intel/", "/api/pipeline/", "/api/constructor_profiles",
        "/api/mclaren/", "/api/victory/", "/api/anomaly/", "/api/forecast/",
        "/api/mccar-telemetry/", "/api/mccar-race-telemetry/",
        "/api/mccar-summary", "/api/mcdriver-summary",
        "/api/f1data/", "/api/mccar/", "/api/mcdriver/", "/api/mcracecontext/",
        "/api/strategy/", "/api/team_intel/",
    )
    async def dispatch(self, request: _Request, call_next):
        path = request.scope["path"]
        for prefix in self._PREFIXES:
            if path.startswith(prefix):
                request.scope["path"] = path.replace("/api/", "/api/local/", 1)
                break
        return await call_next(request)

app.add_middleware(_RewriteMiddleware)

# Mount 3D model generation routes
app.include_router(model_3d_router)
mount_3d_static(app)

# Mount OmniSuite routers
app.include_router(omni_health_router)
app.include_router(omni_analytics_router)
app.include_router(omni_rag_router)
app.include_router(omni_kex_router)
app.include_router(omni_doc_router)
app.include_router(omni_data_router)
app.include_router(omni_bedding_router)
app.include_router(omni_vis_router)
app.include_router(omni_dapt_router)
app.include_router(opponents_router)
app.include_router(updater_router)
app.include_router(advantage_router)
app.include_router(omni_agents_router)
app.include_router(radio_router)
app.include_router(aim_router)

# Lazy-init singletons
_groq: Groq | None = None
_vs = None
_embedder: NomicEmbedder | None = None
_clip_index: dict | None = None
_clip_embedder = None

CLIP_INDEX_PATH = Path(__file__).parent.parent / "f1data" / "McMedia" / "clip_index.json"


def get_groq() -> Groq:
    global _groq
    if _groq is None:
        _groq = Groq(api_key=os.getenv("GROQ_API_KEY"))
    return _groq


def get_vs():
    global _vs
    if _vs is None:
        _vs = get_vector_store()
    return _vs


def get_embedder() -> NomicEmbedder:
    global _embedder
    if _embedder is None:
        _embedder = NomicEmbedder()
    return _embedder


def get_clip_index() -> dict:
    """Load pre-built CLIP index from MongoDB (preferred) or disk fallback."""
    global _clip_index
    if _clip_index is None:
        # Try MongoDB first
        try:
            db = get_data_db()
            doc = db["clip_index"].find_one({"_type": "clip_index"}, {"_id": 0, "_type": 0})
            if doc and "images" in doc:
                _clip_index = doc
        except Exception:
            pass
        # Fall back to local file
        if _clip_index is None:
            if not CLIP_INDEX_PATH.exists():
                raise FileNotFoundError(
                    "CLIP index not found in MongoDB or on disk. "
                    "Run: python pipeline/push_clip_index_to_mongo.py"
                )
            with open(CLIP_INDEX_PATH) as f:
                _clip_index = json.load(f)
        # Pre-compute numpy arrays for fast search
        _clip_index["_image_vecs"] = np.array(
            [img["embedding"] for img in _clip_index["images"]]
        )
    return _clip_index


def get_clip_embedder():
    """Lazy-load CLIP embedder for query embedding."""
    global _clip_embedder
    if _clip_embedder is None:
        from pipeline.embeddings import CLIPEmbedder
        _clip_embedder = CLIPEmbedder()
    return _clip_embedder


# ── Models ───────────────────────────────────────────────────────────────

class ChatMessage(BaseModel):
    role: str  # "user" or "assistant"
    content: str


class ChatRequest(BaseModel):
    message: str
    history: list[ChatMessage] = []
    data_types: list[str] | None = None


class ChatResponse(BaseModel):
    answer: str
    sources: list[dict]


# ── System Prompt ────────────────────────────────────────────────────────

SYSTEM_PROMPT = """You are the F1 OmniSense Knowledge Agent — an expert on Formula 1 technical regulations, car specifications, equipment, and engineering standards.

You have access to a knowledge base extracted from FIA 2024 Technical Regulations and related engineering documents. When answering questions:

1. Use ONLY the provided context to answer. If the context doesn't contain enough information, say so clearly.
2. Cite specific regulation IDs (e.g., "Article 3.5.2") and page numbers when available.
3. Be precise with numerical values, units, and tolerances.
4. For dimensional specifications, always include the value and unit.
5. When discussing equipment, reference tags and types.
6. Keep answers concise but thorough. Use bullet points for lists.
7. If a question is ambiguous, briefly clarify what you're answering.

You are speaking with F1 engineers and technical staff — use appropriate technical language."""


# ── RAG Pipeline ─────────────────────────────────────────────────────────

def _text_search_fallback(query: str, k: int = 8, data_types: list[str] | None = None) -> list[dict]:
    """Fallback: keyword search on f1_knowledge when vector search is unavailable."""
    vs = get_vs()
    coll = vs.collection
    keywords = [w for w in query.split() if len(w) > 2]
    if not keywords:
        keywords = query.split()
    regex = "|".join(keywords)
    try:
        match_filter: dict = {"page_content": {"$regex": regex, "$options": "i"}}
        if data_types:
            match_filter["metadata.data_type"] = {"$in": data_types}
        results = list(
            coll.find(
                match_filter,
                {"page_content": 1, "metadata": 1, "_id": 0},
            ).limit(k)
        )
    except Exception:
        results = list(coll.find({}, {"page_content": 1, "metadata": 1, "_id": 0}).limit(k))
    sources = []
    for r in results:
        meta = r.get("metadata", {})
        sources.append({
            "content": r.get("page_content", ""),
            "data_type": meta.get("data_type", ""),
            "category": meta.get("category", ""),
            "source": meta.get("source", ""),
            "page": meta.get("page", 0),
        })
    return sources


def retrieve_context(query: str, k: int = 8, data_types: list[str] | None = None) -> list[dict]:
    """Retrieve relevant documents using text search (fast) with optional
    vector search upgrade if embedder is already loaded."""
    global _embedder
    vs_filter = {"metadata.data_type": {"$in": data_types}} if data_types else None
    # If embedder is already loaded, use vector search
    if _embedder is not None:
        try:
            vs = get_vs()
            query_vec = _embedder.embed_query(query)
            docs = vs.similarity_search(query, k=k, query_embedding=query_vec, filter=vs_filter)
            sources = []
            for doc in docs:
                sources.append({
                    "content": doc.page_content,
                    "data_type": doc.metadata.get("data_type", ""),
                    "category": doc.metadata.get("category", ""),
                    "source": doc.metadata.get("source", ""),
                    "page": doc.metadata.get("page", 0),
                })
            return sources
        except Exception as e:
            print(f"  Vector search failed ({e}), falling back to text search")
    # Fast text search — no model loading required
    return _text_search_fallback(query, k, data_types=data_types)


def build_rag_prompt(query: str, sources: list[dict], history: list[ChatMessage]) -> list[dict]:
    """Build the full prompt with system, context, history, and user query."""
    context_parts = []
    for i, src in enumerate(sources, 1):
        context_parts.append(
            f"[{i}] ({src['data_type']}/{src['category']}) "
            f"Page {src['page']} — {src['source']}\n{src['content']}"
        )
    context_block = "\n\n".join(context_parts)

    messages = [{"role": "system", "content": SYSTEM_PROMPT}]

    # Add conversation history (last 10 messages)
    for msg in history[-10:]:
        messages.append({"role": msg.role, "content": msg.content})

    # User message with context
    user_prompt = f"""CONTEXT FROM KNOWLEDGE BASE:
{context_block}

USER QUESTION:
{query}

Answer based on the context above. Cite regulation IDs and page numbers where applicable."""

    messages.append({"role": "user", "content": user_prompt})
    return messages


# ── Endpoints ────────────────────────────────────────────────────────────

@app.post("/api/fleet/diagnose")
@app.post("/fleet/diagnose")
async def fleet_diagnose_proxy(request: Request):
    """Proxy Gen UI requests to Vite dev server if running, else fallback to plain chat."""
    import httpx
    body = await request.body()
    try:
        async with httpx.AsyncClient(timeout=60) as client:
            resp = await client.post(
                f"http://127.0.0.1:{os.getenv('VITE_PORT', '8889')}/api/fleet/diagnose",
                content=body,
                headers={"content-type": "application/json"},
            )
        from starlette.responses import Response as _Resp
        return _Resp(
            content=resp.content,
            status_code=resp.status_code,
            media_type=resp.headers.get("content-type", "application/json"),
        )
    except Exception:
        # Vite not running — extract message and fallback to plain chat
        import json as _json
        text = ""
        try:
            data = _json.loads(body)
            msgs = data.get("messages", [])
            last_user = next((m for m in reversed(msgs) if m.get("role") == "user"), None)
            if last_user:
                parts = last_user.get("parts", [])
                text = next((p.get("text", "") for p in parts if p.get("type") == "text"), "")
            if not text:
                text = data.get("message", "")
        except Exception:
            pass
        if not text:
            raise HTTPException(400, "No message found in request")
        return chat(ChatRequest(message=text))


@app.post("/chat", response_model=ChatResponse)
@app.post("/api/chat", response_model=ChatResponse)
def chat(req: ChatRequest):
    """RAG chat endpoint. Delegates to OmniRAG when USE_OMNIRAG=true."""
    if USE_OMNIRAG:
        from pipeline.omni_rag_router import chat as _omni_chat, ChatRequest as _OmniReq
        result = _omni_chat(_OmniReq(message=req.message))
        return ChatResponse(answer=result["answer"], sources=result.get("sources", []))

    # 1. Retrieve context
    sources = retrieve_context(req.message, k=8, data_types=req.data_types)

    # 2. Build prompt
    messages = build_rag_prompt(req.message, sources, req.history)

    # 3. Generate answer via Groq
    groq = get_groq()
    completion = groq.chat.completions.create(
        model=GROQ_MODEL,
        messages=messages,
        temperature=0.3,
        max_tokens=2048,
    )

    answer = completion.choices[0].message.content

    return ChatResponse(answer=answer, sources=sources)


@app.api_route("/health", methods=["GET", "HEAD"])
def health():
    # Lightweight health check — no model loading, just DB connectivity
    try:
        db = get_data_db()
        db.command("ping")
        db_ok = True
    except Exception:
        db_ok = False
    return {
        "status": "ok" if db_ok else "degraded",
        "model": GROQ_MODEL,
        "db": "connected" if db_ok else "error",
        "services": ["knowledge_agent", "3d_model_gen"],
    }


# ── CLIP Visual Search ───────────────────────────────────────────────

@app.get("/visual-search")
def visual_search(q: str, k: int = 8):
    """Search images by text query using CLIP embeddings."""
    index = get_clip_index()
    clip = get_clip_embedder()

    # Embed text query into CLIP space
    query_vec = np.array(clip.embed_text(q))
    query_vec = query_vec / np.linalg.norm(query_vec)

    # Cosine similarity against all image embeddings
    image_vecs = index["_image_vecs"]
    image_norms = image_vecs / np.linalg.norm(image_vecs, axis=1, keepdims=True)
    similarities = query_vec @ image_norms.T

    # Top-k results
    top_indices = np.argsort(similarities)[::-1][:k]

    results = []
    for idx in top_indices:
        img = index["images"][int(idx)]
        results.append({
            "path": img["path"],
            "score": round(float(similarities[idx]), 4),
            "auto_tags": img["auto_tags"],
            "source_video": img["source_video"],
            "frame_index": img["frame_index"],
        })

    return {"query": q, "results": results}


@app.get("/api/visual-tags")
def visual_tags():
    """Return all images with their auto-tags for gallery filtering."""
    try:
        index = get_clip_index()
    except FileNotFoundError:
        return {"images": [], "tags": [], "stats": {}}

    images = []
    for img in index["images"]:
        images.append({
            "path": img["path"],
            "auto_tags": img["auto_tags"],
            "source_video": img["source_video"],
            "frame_index": img["frame_index"],
        })

    # Collect all unique tag labels with their max scores
    tag_summary: dict[str, float] = {}
    for img in index["images"]:
        for tag in img["auto_tags"]:
            label = tag["label"]
            if label not in tag_summary or tag["score"] > tag_summary[label]:
                tag_summary[label] = tag["score"]

    top_tags = sorted(tag_summary.items(), key=lambda x: -x[1])

    return {
        "images": images,
        "tags": [{"label": t[0], "max_score": round(t[1], 4)} for t in top_tags],
        "stats": index["stats"],
    }


# ── Document Upload & Ingestion ──────────────────────────────────────────

ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt", ".csv", ".json", ".md"}


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF using PyMuPDF."""
    import fitz
    doc = fitz.open(file_path)
    pages = []
    for page in doc:
        pages.append(page.get_text())
    doc.close()
    return "\n\n".join(pages)


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX using python-docx."""
    from docx import Document as DocxDocument
    doc = DocxDocument(file_path)
    return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())


def extract_text_from_plain(file_path: str) -> str:
    """Read plain text files (txt, csv, json, md)."""
    return Path(file_path).read_text(encoding="utf-8", errors="replace")


EXTRACTORS = {
    ".pdf": extract_text_from_pdf,
    ".docx": extract_text_from_docx,
    ".txt": extract_text_from_plain,
    ".csv": extract_text_from_plain,
    ".json": extract_text_from_plain,
    ".md": extract_text_from_plain,
}


@app.post("/upload")
@app.post("/api/upload")
async def upload_document(file: UploadFile = File(...)):
    """Upload a document, extract text, embed, and ingest into the knowledge base.

    Tries OmniDoc enhanced processing first (richer metadata, table extraction),
    falls back to legacy extractor if OmniDoc is unavailable.
    """
    filename = file.filename or "unknown"
    ext = Path(filename).suffix.lower()

    # Try OmniDoc route first (supports more formats, extracts tables/tags)
    try:
        from pipeline.omni_doc_router import upload_and_ingest
        # Reset file position and delegate
        await file.seek(0)
        return await upload_and_ingest(file=file, deep_extract=False, chunk_size=1000, chunk_overlap=200)
    except ImportError:
        pass
    except Exception as e:
        logger.warning("OmniDoc upload failed (%s), falling back to legacy", e)

    # Legacy fallback
    await file.seek(0)

    if ext not in ALLOWED_EXTENSIONS:
        return {"filename": filename, "status": "error",
                "error": f"Unsupported format: {ext}. Allowed: {', '.join(sorted(ALLOWED_EXTENSIONS))}"}

    # Save to temp file
    with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
        content = await file.read()
        tmp.write(content)
        tmp_path = tmp.name

    try:
        # Extract text
        extractor = EXTRACTORS[ext]
        text = extractor(tmp_path)

        if not text.strip():
            return {"filename": filename, "status": "error", "error": "No text extracted from file"}

        # Chunk text
        from langchain_text_splitters import RecursiveCharacterTextSplitter
        splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
        chunks = splitter.split_text(text)

        if not chunks:
            return {"filename": filename, "status": "error", "error": "No chunks created"}

        # Build LangChain Documents
        from langchain_core.documents import Document
        docs = []
        for i, chunk in enumerate(chunks):
            docs.append(Document(
                page_content=chunk,
                metadata={
                    "data_type": "uploaded_document",
                    "category": "user_upload",
                    "source": filename,
                    "chunk": i + 1,
                    "total_chunks": len(chunks),
                },
            ))

        # Embed
        embedder = get_embedder()
        texts = [doc.page_content for doc in docs]
        embeddings: list[list[float]] = []
        for i in range(0, len(texts), 32):
            batch = texts[i:i + 32]
            embeddings.extend(embedder.embed(batch))

        # Upsert to Atlas
        vs = get_vs()
        count = vs.upsert_documents(docs, embeddings)

        return {
            "filename": filename,
            "status": "ok",
            "chunks": count,
            "text_length": len(text),
        }

    except Exception as e:
        return {"filename": filename, "status": "error", "error": str(e)}
    finally:
        os.unlink(tmp_path)


# ── Data API — serve MongoDB collections to frontend ─────────────────────

from pymongo import MongoClient as _MongoClient

_data_client: _MongoClient | None = None
_data_db = None

def get_data_db():
    global _data_client, _data_db
    if _data_db is None:
        _data_client = _MongoClient(os.getenv("MONGODB_URI", ""))
        _data_db = _data_client[os.getenv("MONGODB_DB", "marip_f1")]
        # Share this connection with the opponents profiler
        init_profiler_with_db(_data_db)
        # Ensure feature_store index for cross-model caching
        try:
            from omnianalytics.feature_store import ensure_indexes
            ensure_indexes(_data_db)
        except Exception:
            pass
    return _data_db

import base64 as _b64
from fastapi.responses import Response as _Response, FileResponse as _FileResponse
from pathlib import Path as _Path

_MEDIA_DIR = _Path(__file__).resolve().parent.parent / "f1data" / "McMedia"

@app.get("/media/{filename:path}")
async def serve_media(filename: str, request: Request):
    """Serve media files: videos from disk, images from MongoDB."""
    # If path has a slash, it's a subfolder image (e.g., gdino_results/img.jpg)
    if "/" in filename:
        db = get_data_db()
        doc = db["media_frames"].find_one({"path": filename}, {"data_b64": 1, "content_type": 1})
        if doc:
            return _Response(content=_b64.b64decode(doc["data_b64"]), media_type=doc.get("content_type", "image/jpeg"))
        return _Response(status_code=404, content=b"Not found")
    # Single filename — try disk first, then GridFS
    fpath = _MEDIA_DIR / filename
    if fpath.is_file():
        suffix = fpath.suffix.lower()
        mime = {".mp4": "video/mp4", ".webm": "video/webm", ".avi": "video/x-msvideo",
                ".mov": "video/quicktime", ".mkv": "video/x-matroska"}.get(suffix, "application/octet-stream")
        return _FileResponse(fpath, media_type=mime)
    # Fallback: serve from MongoDB GridFS with Range support for video playback
    import gridfs as _gridfs
    db = get_data_db()
    fs = _gridfs.GridFS(db, collection="media_files")
    gf = fs.find_one({"filename": filename})
    if gf:
        total = gf.length
        mime = gf.content_type or "video/mp4"
        range_header = request.headers.get("range")
        if range_header:
            # Parse Range: bytes=start-end
            import re as _re
            m = _re.match(r"bytes=(\d+)-(\d*)", range_header)
            start = int(m.group(1)) if m else 0
            end = int(m.group(2)) if m and m.group(2) else total - 1
            end = min(end, total - 1)
            length = end - start + 1
            gf.seek(start)
            data = gf.read(length)
            return _Response(
                content=data, status_code=206, media_type=mime,
                headers={
                    "Content-Range": f"bytes {start}-{end}/{total}",
                    "Accept-Ranges": "bytes",
                    "Content-Length": str(length),
                },
            )
        # No range — serve full file
        return _Response(
            content=gf.read(), media_type=mime,
            headers={"Accept-Ranges": "bytes", "Content-Length": str(total)},
        )
    return _Response(status_code=404, content=b"Not found")

@app.get("/api/local/jolpica/race_results")
async def jolpica_race_results():
    db = get_data_db()
    docs = list(db["jolpica_race_results"].find({}, {"_id": 0, "ingested_at": 0}))
    return docs

@app.get("/api/local/jolpica/driver_standings")
async def jolpica_driver_standings():
    db = get_data_db()
    docs = list(db["jolpica_driver_standings"].find({}, {"_id": 0, "ingested_at": 0}))
    return docs

@app.get("/api/local/jolpica/constructor_standings")
async def jolpica_constructor_standings():
    db = get_data_db()
    docs = list(db["jolpica_constructor_standings"].find({}, {"_id": 0, "ingested_at": 0}))
    return docs

@app.get("/api/local/constructor_profiles")
async def constructor_profiles(
    season: int | None = None,
    constructor_id: str | None = None,
):
    """Team performance profiles aggregated from race results, qualifying, pit stops, and telemetry."""
    db = get_data_db()
    filt: dict = {}
    if season:
        filt["season"] = season
    if constructor_id:
        filt["constructor_id"] = constructor_id
    docs = list(db["constructor_profiles"].find(filt, {"_id": 0, "updated_at": 0}))
    docs.sort(key=lambda d: (d.get("season", 0), d.get("championship_position") or 99))
    return docs

@app.get("/api/local/constructor_profiles/{constructor_id}")
async def constructor_profile_detail(constructor_id: str, season: int | None = None):
    """Single constructor profile, optionally filtered by season."""
    db = get_data_db()
    filt: dict = {"constructor_id": constructor_id}
    if season:
        filt["season"] = season
    docs = list(db["constructor_profiles"].find(filt, {"_id": 0, "updated_at": 0}))
    if not docs:
        return {"error": f"No profile found for {constructor_id}"}
    docs.sort(key=lambda d: d.get("season", 0))
    return docs[0] if season else docs

@app.get("/api/local/jolpica/qualifying")
async def jolpica_qualifying(season: int | None = None, driver: str | None = None):
    db = get_data_db()
    filt: dict = {}
    if season:
        filt["season"] = season
    if driver:
        filt["driver_code"] = driver.upper()
    return list(db["jolpica_qualifying"].find(filt, {"_id": 0}).limit(5000))

@app.get("/api/local/jolpica/circuits")
async def jolpica_circuits():
    db = get_data_db()
    pipeline = [
        {"$group": {"_id": "$circuit_id", "race_name": {"$first": "$race_name"}, "races": {"$sum": 1}}},
        {"$sort": {"_id": 1}},
    ]
    return list(db["jolpica_race_results"].aggregate(pipeline))

@app.get("/api/local/jolpica/pit_stops")
async def jolpica_pit_stops(season: int | None = None, circuit: str | None = None):
    db = get_data_db()
    filt: dict = {}
    if season:
        filt["season"] = season
    if circuit:
        filt["circuit_id"] = circuit
    return list(db["jolpica_pit_stops"].find(filt, {"_id": 0}).limit(6000))

@app.get("/api/local/jolpica/sprint_results")
async def jolpica_sprint_results(season: int | None = None):
    db = get_data_db()
    filt: dict = {}
    if season:
        filt["season"] = season
    return list(db["jolpica_sprint_results"].find(filt, {"_id": 0}))

@app.get("/api/local/jolpica/lap_times")
async def jolpica_lap_times():
    return []

@app.get("/api/local/jolpica/drivers")
async def jolpica_drivers():
    db = get_data_db()
    pipeline = [
        {"$group": {"_id": "$driver_id", "code": {"$first": "$driver_code"}, "constructor": {"$last": "$constructor_id"}, "races": {"$sum": 1}}},
        {"$sort": {"races": -1}},
    ]
    return list(db["jolpica_race_results"].aggregate(pipeline))

@app.get("/api/local/jolpica/seasons")
async def jolpica_seasons():
    db = get_data_db()
    return sorted(db["jolpica_race_results"].distinct("season"))

# ── Driver Intelligence endpoints ─────────────────────────────────────

@app.get("/api/local/driver_intel/performance_markers")
async def driver_intel_performance_markers(driver: str | None = None):
    db = get_data_db()
    filt: dict = {}
    if driver:
        filt["Driver"] = driver.upper()
    return list(db["driver_performance_markers"].find(filt, {"_id": 0}))

@app.get("/api/local/driver_intel/overtake_profiles")
async def driver_intel_overtake_profiles(driver: str | None = None):
    db = get_data_db()
    filt: dict = {}
    if driver:
        filt["driver_code"] = driver.upper()
    return list(db["driver_overtake_profiles"].find(filt, {"_id": 0}))

@app.get("/api/local/driver_intel/telemetry_profiles")
async def driver_intel_telemetry_profiles(driver: str | None = None):
    db = get_data_db()
    filt: dict = {}
    if driver:
        filt["driver_code"] = driver.upper()
    return list(db["driver_telemetry_profiles"].find(filt, {"_id": 0}))

@app.get("/api/local/driver_intel/similar/{driver_code}")
async def driver_intel_similar(driver_code: str, k: int = 5, season: int | None = None):
    """Find k most similar drivers using VectorProfiles embeddings."""
    from pipeline.build_vector_profiles import find_similar
    db = get_data_db()
    try:
        results = find_similar(driver_code, k=k, db=db, season=season)
    except ValueError as e:
        raise HTTPException(404, str(e))
    return results


@app.get("/api/local/driver_intel/vector_profile/{driver_code}")
async def driver_intel_vector_profile(driver_code: str, season: int | None = None):
    """Get the full VectorProfile for a driver (without embedding)."""
    db = get_data_db()
    query = {"driver_code": driver_code.upper()}
    if season is not None and season >= 2025:
        query["season"] = season
    else:
        query["season"] = None
    doc = db["VectorProfiles"].find_one(query, {"_id": 0, "embedding": 0})
    if not doc:
        raise HTTPException(404, f"No VectorProfile for {driver_code}")
    return doc


@app.get("/api/local/team_intel/similar/{team_name}")
async def team_intel_similar(team_name: str, k: int = 5, season: int | None = None):
    """Find k most similar teams using averaged VectorProfiles embeddings."""
    from pipeline.build_vector_profiles import find_similar_teams
    db = get_data_db()
    try:
        results = find_similar_teams(team_name, k=k, db=db, season=season)
    except ValueError as e:
        raise HTTPException(404, str(e))
    return results


@app.get("/api/local/team_intel/intra/{team_name}")
async def team_intel_intra(team_name: str, season: int | None = None):
    """Get pairwise driver similarity within a team."""
    from pipeline.build_vector_profiles import get_intra_team_similarity
    db = get_data_db()
    return get_intra_team_similarity(team_name, db=db, season=season)


# ── VictoryProfiles endpoints ────────────────────────────────────────────


@app.get("/api/local/victory/team/{team}/{season}")
async def victory_team_kb(team: str, season: int):
    """Full team KB with driver + car profiles."""
    db = get_data_db()
    kb = db["victory_team_kb"].find_one(
        {"team": {"$regex": f"^{team}$", "$options": "i"}, "season": season},
        {"_id": 0, "embedding": 0},
    )
    if not kb:
        raise HTTPException(404, f"No VictoryProfile for {team} {season}")

    drivers = list(db["victory_driver_profiles"].find(
        {"team": {"$regex": f"^{team}$", "$options": "i"}, "season": season},
        {"_id": 0, "embedding": 0},
    ))
    kb["driver_profiles"] = drivers

    car = db["victory_car_profiles"].find_one(
        {"team": {"$regex": f"^{team}$", "$options": "i"}, "season": season},
        {"_id": 0, "embedding": 0},
    )
    kb["car_profile"] = car

    strategies = list(db["victory_strategy_profiles"].find(
        {"team": {"$regex": f"^{team}$", "$options": "i"}, "season": season},
        {"_id": 0, "embedding": 0},
    ))
    kb["strategy_profiles"] = strategies
    return kb


@app.post("/api/local/victory/compare")
async def victory_compare(body: dict):
    """Compare 2+ teams by cosine similarity + structured diff.

    Body: {"teams": ["McLaren", "Red Bull Racing"], "season": 2024}
    """
    teams = body.get("teams", [])
    season = body.get("season")
    if len(teams) < 2:
        raise HTTPException(400, "Provide at least 2 teams to compare")

    db = get_data_db()
    team_data = {}
    for team in teams:
        doc = db["victory_team_kb"].find_one(
            {"team": {"$regex": f"^{team}$", "$options": "i"}, "season": season},
            {"_id": 0},
        )
        if doc:
            team_data[doc["team"]] = doc

    if len(team_data) < 2:
        raise HTTPException(404, f"Need 2+ teams with profiles. Found: {list(team_data.keys())}")

    team_names = list(team_data.keys())
    similarities = []
    for i in range(len(team_names)):
        for j in range(i + 1, len(team_names)):
            v1 = np.array(team_data[team_names[i]]["embedding"])
            v2 = np.array(team_data[team_names[j]]["embedding"])
            score = float(np.dot(v1, v2) / (np.linalg.norm(v1) * np.linalg.norm(v2) + 1e-10))
            similarities.append({
                "team_a": team_names[i], "team_b": team_names[j],
                "similarity": round(score, 4),
            })

    diffs = []
    for i in range(len(team_names)):
        for j in range(i + 1, len(team_names)):
            meta_a = team_data[team_names[i]].get("metadata", {})
            meta_b = team_data[team_names[j]].get("metadata", {})
            diffs.append(_compute_victory_diff(team_names[i], meta_a, team_names[j], meta_b))

    return {"teams": team_names, "season": season, "similarities": similarities, "diffs": diffs}


@app.post("/api/local/victory/search")
async def victory_search(body: dict):
    """Semantic search across team KBs.

    Body: {"query": "teams with strong brakes", "season": 2024, "k": 5}
    """
    query = body.get("query", "")
    season = body.get("season")
    k = body.get("k", 5)
    if not query:
        raise HTTPException(400, "Provide a query string")

    db = get_data_db()
    from pipeline.embeddings import NomicEmbedder
    embedder = NomicEmbedder()
    query_vec = np.array(embedder.embed_query(query))

    s_filter = {"season": season} if season is not None else {}
    results = []
    for doc in db["victory_team_kb"].find(s_filter, {"_id": 0}):
        if "embedding" not in doc:
            continue
        vec = np.array(doc["embedding"])
        score = float(np.dot(query_vec, vec) / (np.linalg.norm(query_vec) * np.linalg.norm(vec) + 1e-10))
        results.append({
            "team": doc["team"], "season": doc.get("season"),
            "score": round(score, 4),
            "narrative": doc.get("narrative", "")[:300],
            "metadata": doc.get("metadata"),
        })

    results.sort(key=lambda x: x["score"], reverse=True)
    return {"query": query, "results": results[:k]}


@app.post("/api/local/victory/build")
async def victory_build(body: dict = {}):
    """Trigger VictoryProfiles + VectorProfiles rebuild from the frontend."""
    import asyncio
    season = body.get("season")
    rebuild = body.get("rebuild", False)

    results = {}

    # Build VictoryProfiles (4-layer KB)
    try:
        from pipeline.build_victory_profiles import main as build_victory
        await asyncio.to_thread(build_victory, season=season, rebuild=rebuild)
        db = get_data_db()
        results["victory"] = {
            "driver_profiles": db["victory_driver_profiles"].count_documents({}),
            "car_profiles": db["victory_car_profiles"].count_documents({}),
            "strategy_profiles": db["victory_strategy_profiles"].count_documents({}),
            "team_kb": db["victory_team_kb"].count_documents({}),
        }
    except Exception as e:
        logger.exception("VictoryProfiles build failed")
        results["victory"] = {"error": str(e)}

    # Build VectorProfiles (merged + embeddings)
    try:
        from pipeline.build_vector_profiles import main as build_vectors
        await asyncio.to_thread(build_vectors, rebuild=rebuild)
        db = get_data_db()
        results["vector"] = {
            "count": db["VectorProfiles"].count_documents({}),
            "with_embeddings": db["VectorProfiles"].count_documents({"embedding": {"$exists": True}}),
        }
    except Exception as e:
        logger.exception("VectorProfiles build failed")
        results["vector"] = {"error": str(e)}

    return {"status": "complete", **results}


@app.get("/api/local/victory/regression/{team}")
async def victory_regression(team: str, season_a: int = 2023, season_b: int = 2024):
    """Season-over-season diff for internal improvement analysis."""
    db = get_data_db()

    kb_a = db["victory_team_kb"].find_one(
        {"team": {"$regex": f"^{team}$", "$options": "i"}, "season": season_a},
        {"_id": 0, "embedding": 0},
    )
    kb_b = db["victory_team_kb"].find_one(
        {"team": {"$regex": f"^{team}$", "$options": "i"}, "season": season_b},
        {"_id": 0, "embedding": 0},
    )

    if not kb_a or not kb_b:
        missing = []
        if not kb_a:
            missing.append(str(season_a))
        if not kb_b:
            missing.append(str(season_b))
        raise HTTPException(404, f"Missing {team} profile for season(s): {', '.join(missing)}")

    diff = _compute_victory_diff(
        f"{team} {season_a}", kb_a.get("metadata", {}),
        f"{team} {season_b}", kb_b.get("metadata", {}),
    )
    return {
        "team": team, "season_a": season_a, "season_b": season_b,
        "diff": diff,
        "narrative_a": kb_a.get("narrative", ""),
        "narrative_b": kb_b.get("narrative", ""),
    }


def _compute_victory_diff(name_a: str, meta_a: dict, name_b: str, meta_b: dict) -> dict:
    """Compute structured diff between two team metadata blobs."""
    diff = {"teams": [name_a, name_b], "car": {}, "drivers": {}}

    car_a = meta_a.get("car", {})
    car_b = meta_b.get("car", {})

    # Health systems diff
    sys_a = (car_a.get("health") or {}).get("systems", {})
    sys_b = (car_b.get("health") or {}).get("systems", {})
    all_systems = sorted(set(list(sys_a.keys()) + list(sys_b.keys())))
    system_diffs = []
    for sys_name in all_systems:
        va, vb = sys_a.get(sys_name), sys_b.get(sys_name)
        if va is not None and vb is not None:
            system_diffs.append({"system": sys_name, name_a: va, name_b: vb, "delta": round(vb - va, 1)})
    diff["car"]["systems"] = system_diffs

    # Constructor stats diff
    con_a = car_a.get("constructor") or {}
    con_b = car_b.get("constructor") or {}
    con_diffs = []
    for field in ("total_wins", "total_podiums", "dnf_rate", "avg_finish_position",
                  "avg_pit_duration_s", "q3_rate"):
        va, vb = con_a.get(field), con_b.get(field)
        if va is not None and vb is not None:
            delta = round(vb - va, 3) if isinstance(vb, float) else vb - va
            con_diffs.append({"metric": field, name_a: va, name_b: vb, "delta": delta})
    diff["car"]["constructor"] = con_diffs

    # Strategy diff
    strat_a = meta_a.get("strategy", {})
    strat_b = meta_b.get("strategy", {})
    strat_diffs = []
    for field in ("team_undercut_aggression", "team_avg_tyre_life", "team_one_stop_freq"):
        va, vb = strat_a.get(field), strat_b.get(field)
        if va is not None and vb is not None:
            strat_diffs.append({"metric": field, name_a: va, name_b: vb, "delta": round(vb - va, 3)})
    diff["strategy"] = strat_diffs

    return diff


def _extract_summary(text: str, max_chars: int = 300) -> str:
    """Extract first meaningful paragraph as summary from WISE output."""
    import re
    # Strip markdown headers and numbered sections
    cleaned = re.sub(r'^#{1,4}\s+\d*\.?\s*.*$', '', text, flags=re.MULTILINE).strip()
    # Split into paragraphs
    paragraphs = [p.strip() for p in cleaned.split('\n\n') if p.strip() and len(p.strip()) > 30]
    if not paragraphs:
        return text[:max_chars].rsplit(' ', 1)[0] + '...' if len(text) > max_chars else text
    summary = paragraphs[0]
    if len(summary) > max_chars:
        summary = summary[:max_chars].rsplit(' ', 1)[0] + '...'
    return summary


def _scores_from_fingerprint(insight) -> dict:
    """Build radar scores from WISE fingerprint numeric_stats (100% grounded).

    Picks up to 7 columns with the highest coefficient of variation (most
    interesting spread), normalizes each mean to 0-10 within its own min-max range.
    """
    fp = getattr(insight, "fingerprint", None)
    if not fp:
        return {}
    stats = getattr(fp, "numeric_stats", None) or (fp.get("numeric_stats") if isinstance(fp, dict) else None)
    if not stats:
        return {}

    # Skip identifier / non-metric columns
    skip = {"_id", "updated_at", "driver_code", "Driver", "season", "year",
            "Year", "Race", "round", "samples", "laps"}

    # Friendly display names for known columns
    labels = {
        # telemetry_race_summary columns
        "avg_speed": "Avg Speed",
        "top_speed": "Top Speed",
        "avg_rpm": "Engine RPM",
        "max_rpm": "Peak RPM",
        "avg_throttle": "Throttle %",
        "brake_pct": "Braking",
        "drs_pct": "DRS Usage",
        # jolpica_race_results columns
        "grid": "Grid Position",
        "position": "Race Finish",
        "points": "Points",
        "positions_gained": "Overtaking",
        # aggregate profile columns (legacy)
        "avg_race_speed_kmh": "Race Speed",
        "lap_time_consistency_std": "Consistency",
        "degradation_slope_s_per_lap": "Tyre Deg",
        "overtake_ratio": "Overtake Rate",
        "late_race_delta_s": "Late Race Pace",
        "avg_braking_g": "Braking Force",
        "full_throttle_pct": "Throttle %",
        "drs_gain_kmh": "DRS Gain",
        "braking_consistency_std": "Brake Consistency",
        "avg_stint_laps": "Stint Length",
        "brake_overlap_throttle_pct": "Brake Overlap",
        # anomaly columns
        "overallHealth": "Overall Health",
        # circuit columns
        "estimated_corners": "Corner Count",
        "drs_zones": "DRS Zones",
        "elevation_gain_m": "Elevation",
        "est_pit_lane_loss_s": "Pit Loss",
        "avg_temp_c": "Temperature",
        "air_density_kg_m3": "Air Density",
        "downforce_loss_pct": "DF Loss",
        "pole_win_rate": "Pole-Win %",
        "avg_positions_gained": "Overtaking",
        "dnf_rate": "DNF Rate",
        "pole_converted": "Pole Conv.",
        "dnf_count": "DNF Count",
        "entries": "Race Entries",
    }

    # Score each column by coefficient of variation (interestingness)
    candidates = []
    for col, s in stats.items():
        if col in skip:
            continue
        mean = s.get("mean", 0)
        std = s.get("std", 0)
        mn = s.get("min", 0)
        mx = s.get("max", 0)
        rng = mx - mn
        if rng == 0:
            continue
        cv = abs(std / mean) if mean != 0 else std
        # Normalize mean to 0-10 within column range
        norm = round((mean - mn) / rng * 10, 1)
        label = labels.get(col, col.replace("_", " ").title()[:18])
        candidates.append((cv, label, max(0, min(10, norm))))

    # Pick top 7 by interestingness, sort alphabetically for display
    candidates.sort(key=lambda x: -x[0])
    top = sorted(candidates[:7], key=lambda x: x[1])
    return {label: val for _, label, val in top}


@app.post("/api/local/driver_intel/kex/compare")
async def driver_compare_kex(body: dict):
    """KeX-powered driver comparison — Gen UI with WISE fingerprint scores."""
    import time as _time
    import hashlib, json as _json
    import pandas as pd

    drivers = body.get("drivers", [])
    if len(drivers) < 2 or len(drivers) > 4:
        raise HTTPException(400, "Provide 2-4 driver codes")

    codes = [d.upper() for d in drivers]
    db = get_data_db()
    coll = db["kex_comparison_briefings"]

    # ── Fetch opponent profiles ───────────────────────────────────────
    cache_key = "+".join(sorted(codes))
    rows = []
    for code in codes:
        doc = db["opponent_profiles"].find_one({"driver_code": code}, {"_id": 0})
        if doc:
            rows.append(doc)
    if len(rows) < 2:
        raise HTTPException(404, f"Need opponent_profiles for at least 2 drivers, found {len(rows)}")

    # ── Check cache (hash includes actual data content) ────────────────
    hash_payload = "v2:" + _json.dumps(
        {d.get("driver_code", ""): d.get("updated_at", "") for d in rows},
        sort_keys=True, default=str,
    )
    data_hash = hashlib.sha256(hash_payload.encode()).hexdigest()[:16]

    existing = coll.find_one({"cache_key": cache_key})
    if existing and existing.get("data_hash") == data_hash:
        existing.pop("_id", None)
        return existing

    df = pd.DataFrame(rows)
    # Keep driver_code for reference, drop non-numeric metadata
    meta_cols = ["driver_id", "forename", "surname", "nationality", "dob",
                 "seasons", "updated_at", "career_stats_updated_at",
                 "jolpica_enriched_at", "driver_number", "age", "preferred_compound"]
    df = df.drop(columns=[c for c in meta_cols if c in df.columns], errors="ignore")

    # ── WISE extraction ────────────────────────────────────────────────
    from omnikex.pillars import extract_realtime
    from omnikex._types import KexLLMConfig, LLMProvider

    driver_list = ", ".join(codes)
    persona = (
        f"You are an F1 strategist comparing {len(codes)} drivers head-to-head. "
        "Focus on what differentiates them — where their strengths and weaknesses diverge. "
        "Highlight strategic matchup implications: who has the edge in qualifying, race pace, "
        "tyre management, late-race performance, and overtaking. "
        "IMPORTANT: Write a concise flowing comparison in 3-4 paragraphs. "
        "Do NOT use numbered sections, headers, or the WISE framework structure in your output. "
        "Do NOT include markdown tables. Just write clear engineering prose."
    )
    query = f"Compare these F1 drivers head-to-head: {driver_list}. Analyze their relative strengths, weaknesses, and strategic matchup implications."

    text = model_used = provider_used = None
    grounding_score = None
    try:
        llm_cfg = KexLLMConfig(
            provider=LLMProvider.GROQ,
            model="llama-3.3-70b-versatile",
            task_type="realtime",
            temperature=0.15,
        )
        insight = extract_realtime(df, query, llm_config=llm_cfg, persona_context=persona, response_length="medium", verify=True)
        text = insight.text
        model_used = insight.model_used
        provider_used = insight.provider_used
        grounding_score = insight.grounding.grounding_score if insight.grounding else None
    except Exception as e:
        logger.warning("Compare KeX Groq failed, falling back to edge: %s", e)

    if not text:
        try:
            llm_cfg = KexLLMConfig(
                provider=LLMProvider.OLLAMA,
                model="qwen3.5:9b",
                task_type="realtime",
                temperature=0.15,
            )
            insight = extract_realtime(df, query, llm_config=llm_cfg, persona_context=persona, response_length="medium", verify=False)
            text = insight.text
            model_used = insight.model_used
            provider_used = insight.provider_used
            grounding_score = None
        except Exception as e2:
            logger.exception("Compare KeX edge fallback also failed")
            raise HTTPException(500, f"WISE extraction failed: {e2}")

    scores = _scores_from_fingerprint(insight)

    now = _time.time()
    result = {
        "cache_key": cache_key,
        "drivers": codes,
        "text": text,
        "model_used": model_used,
        "provider_used": provider_used,
        "scores": scores,
        "summary": _extract_summary(text),
        "grounding_score": grounding_score,
        "generated_at": now,
        "data_hash": data_hash,
    }

    coll.replace_one({"cache_key": cache_key}, result, upsert=True)
    result.pop("_id", None)
    return result


@app.post("/api/local/driver_intel/kex/{driver_code}")
async def driver_intel_kex(driver_code: str, force: bool = False):
    """Generate WISE knowledge extraction briefing for a driver.

    Uses on-screen data (performance markers, overtake profile, telemetry style,
    system health) to create a natural language intelligence briefing.
    Auto-regenerates when source data changes (via data hash).
    """
    import time as _time
    import hashlib, json as _json

    code = driver_code.upper()
    db = get_data_db()
    coll = db["kex_driver_briefings"]

    # ── Gather full historical race data (pre-2025) ──────────────────
    import pandas as pd

    telem_docs = list(db["telemetry_race_summary"].find(
        {"Driver": code}, {"_id": 0, "compounds": 0}
    ))
    results_docs = list(db["jolpica_race_results"].find(
        {"driver_code": code}, {"_id": 0}
    ))

    if not telem_docs and not results_docs:
        raise HTTPException(404, f"No historical data for driver {code}")

    # ── Compute data hash to detect changes (includes content sample) ─
    sample_telem = telem_docs[-1] if telem_docs else {}
    sample_results = results_docs[-1] if results_docs else {}
    hash_payload = "v5:" + _json.dumps(
        {"telem_n": len(telem_docs), "results_n": len(results_docs),
         "last_telem": {k: v for k, v in sample_telem.items() if k not in ("_id", "compounds")},
         "last_result": {k: v for k, v in sample_results.items() if k != "_id"}},
        sort_keys=True, default=str,
    )
    data_hash = hashlib.sha256(hash_payload.encode()).hexdigest()[:16]

    # Return cached if hash matches (data unchanged)
    if not force:
        existing = coll.find_one({"driver_code": code})
        if existing and existing.get("data_hash") == data_hash:
            existing.pop("_id", None)
            return existing

    # ── Build rich DataFrame for WISE ─────────────────────────────────
    telem_df = pd.DataFrame(telem_docs) if telem_docs else pd.DataFrame()
    results_df = pd.DataFrame(results_docs) if results_docs else pd.DataFrame()

    # Merge on (Year/season, Race/round) if both exist
    if not telem_df.empty and not results_df.empty:
        # Standardize join keys
        if "Year" in telem_df.columns and "season" in results_df.columns:
            results_df["Year"] = results_df["season"].astype(int)
            results_df = results_df.drop(columns=["season"])
        if "Race" in telem_df.columns and "round" in results_df.columns:
            # Race is str in telemetry, round is int in results — coerce both to int
            telem_df["Race"] = pd.to_numeric(telem_df["Race"], errors="coerce")
            results_df["Race"] = results_df["round"].astype(int)
            results_df = results_df.drop(columns=["round"])
        merge_on = [c for c in ["Year", "Race"] if c in telem_df.columns and c in results_df.columns]
        if merge_on:
            df = telem_df.merge(results_df, on=merge_on, how="outer", suffixes=("", "_res"))
        else:
            df = pd.concat([telem_df, results_df], ignore_index=True)
    elif not telem_df.empty:
        df = telem_df
    else:
        df = results_df

    # Drop non-numeric clutter columns that confuse fingerprinting
    drop_cols = [c for c in ["driver_code", "Driver", "race_name", "circuit_id",
                              "date", "status", "fastest_lap_time", "driver_code_res"] if c in df.columns]
    df = df.drop(columns=drop_cols, errors="ignore")

    # ── WISE extraction via OmniKeX ──────────────────────────────────
    from omnikex.pillars import extract_realtime
    from omnikex._types import KexLLMConfig, LLMProvider

    persona = (
        "You are an F1 performance analyst preparing a driver intelligence briefing. "
        "Focus on driving style, strengths, weaknesses, race strategy implications, "
        "and competitor context. "
        "IMPORTANT: Write a concise flowing briefing in 3-4 paragraphs. "
        "Do NOT use numbered sections, headers, or the WISE framework structure in your output. "
        "Do NOT include markdown tables. Just write clear engineering prose."
    )
    query = f"Analyze driver {code}'s complete performance profile — driving style, strengths, weaknesses, and strategic implications."

    # Primary: Groq cloud
    text = model_used = provider_used = None
    grounding_score = None
    try:
        llm_cfg = KexLLMConfig(
            provider=LLMProvider.GROQ,
            model="llama-3.3-70b-versatile",
            task_type="realtime",
            temperature=0.15,
        )
        insight = extract_realtime(df, query, llm_config=llm_cfg, persona_context=persona, response_length="medium", verify=True)
        text = insight.text
        model_used = insight.model_used
        provider_used = insight.provider_used
        grounding_score = insight.grounding.grounding_score if insight.grounding else None
    except Exception as e:
        logger.warning("Driver KeX Groq failed for %s, falling back to edge model: %s", code, e)

    # Fallback: Ollama edge model
    if not text:
        try:
            llm_cfg = KexLLMConfig(
                provider=LLMProvider.OLLAMA,
                model="qwen3.5:9b",
                task_type="realtime",
                temperature=0.15,
            )
            insight = extract_realtime(df, query, llm_config=llm_cfg, persona_context=persona, response_length="medium", verify=False)
            text = insight.text
            model_used = insight.model_used
            provider_used = insight.provider_used
            grounding_score = None
        except Exception as e2:
            logger.exception("Driver KeX edge fallback also failed for %s", code)
            raise HTTPException(500, f"WISE extraction failed: {e2}")

    scores = _scores_from_fingerprint(insight)

    now = _time.time()
    result = {
        "driver_code": code,
        "text": text,
        "model_used": model_used,
        "provider_used": provider_used,
        "scores": scores,
        "summary": _extract_summary(text),
        "grounding_score": grounding_score,
        "generated_at": now,
        "data_hash": data_hash,
    }

    coll.replace_one({"driver_code": code}, result, upsert=True)
    result.pop("_id", None)
    return result


# ── Team-level KeX Briefings (declared before driver-level for route priority) ──

@app.post("/api/local/mccar-telemetry/kex/team/{constructor_id}")
async def car_telemetry_kex_team(constructor_id: str, body: dict = {}):
    """Generate a single team-level car telemetry KeX briefing aggregating all drivers."""
    import time as _time
    import hashlib, json as _json
    import pandas as pd

    year = body.get("year")
    if not year:
        raise HTTPException(400, "year is required in request body")

    db = get_data_db()
    coll = db["kex_car_telemetry_briefings"]

    profile = db["constructor_profiles"].find_one(
        {"constructor_id": constructor_id, "season": int(year)}, {"_id": 0}
    )
    if not profile:
        profile = db["constructor_profiles"].find_one(
            {"constructor_id": constructor_id}, {"_id": 0}
        )
    if not profile or not profile.get("drivers"):
        raise HTTPException(404, f"No profile for {constructor_id} in {year}")

    team_name = profile.get("constructor_name", constructor_id)
    driver_codes = [d["driver_code"] for d in profile["drivers"] if d.get("driver_code")]
    if not driver_codes:
        raise HTTPException(404, f"No drivers found for {constructor_id}")

    all_rows = []
    for code in driver_codes:
        docs = list(db["telemetry_race_summary"].find(
            {"Driver": code, "Year": int(year)}, {"_id": 0, "compounds": 0},
        ))
        for doc in docs:
            doc["_driver"] = code
        all_rows.extend(docs)

    if not all_rows:
        raise HTTPException(404, f"No telemetry for {constructor_id} drivers in {year}")

    cache_key = f"team:{constructor_id}:{year}:season"
    hash_payload = f"v1:{cache_key}:{len(all_rows)}"
    last = {k: v for k, v in all_rows[-1].items() if k not in ("_id",)}
    hash_payload += ":" + _json.dumps(last, sort_keys=True, default=str)
    data_hash = hashlib.sha256(hash_payload.encode()).hexdigest()[:16]

    existing = coll.find_one({"cache_key": cache_key})
    if existing and existing.get("data_hash") == data_hash:
        existing.pop("_id", None)
        return existing

    df = pd.DataFrame(all_rows)
    drop_cols = [c for c in ["Driver", "driver_acronym", "meeting_name",
                              "session_name", "session_type", "compound",
                              "_id", "_driver"] if c in df.columns]
    df_clean = df.drop(columns=drop_cols, errors="ignore")

    from omnikex.pillars import extract_realtime
    from omnikex._types import KexLLMConfig, LLMProvider

    drivers_str = ", ".join(driver_codes)
    persona = (
        f"You are an F1 telemetry engineer analyzing {team_name}'s {year} season car performance "
        f"across their full driver lineup ({drivers_str}). "
        "Compare how each driver extracts performance from the car — speed characteristics, "
        "braking behavior, DRS usage, tyre management differences. "
        "Identify team-wide strengths and weaknesses, plus driver-specific divergences. "
        "IMPORTANT: Write a concise flowing team analysis in 4-5 paragraphs. "
        "Do NOT use numbered sections, headers, or the WISE framework structure in your output. "
        "Do NOT include markdown tables. Just write clear engineering prose."
    )
    query = (
        f"Analyze {team_name}'s car telemetry across {year} — team-wide performance patterns, "
        f"driver comparison ({drivers_str}), and areas of strength and concern."
    )

    text = model_used = provider_used = None
    grounding_score = None
    try:
        llm_cfg = KexLLMConfig(provider=LLMProvider.GROQ, model="llama-3.3-70b-versatile",
                                task_type="realtime", temperature=0.15)
        insight = extract_realtime(df_clean, query, llm_config=llm_cfg, persona_context=persona, response_length="medium", verify=True)
        text, model_used, provider_used = insight.text, insight.model_used, insight.provider_used
        grounding_score = insight.grounding.grounding_score if insight.grounding else None
    except Exception as e:
        logger.warning("Team telemetry KeX Groq failed for %s: %s", constructor_id, e)

    if not text:
        try:
            llm_cfg = KexLLMConfig(provider=LLMProvider.OLLAMA, model="qwen3.5:9b",
                                    task_type="realtime", temperature=0.15)
            insight = extract_realtime(df_clean, query, llm_config=llm_cfg, persona_context=persona, response_length="medium", verify=False)
            text, model_used, provider_used = insight.text, insight.model_used, insight.provider_used
            grounding_score = None
        except Exception as e2:
            logger.exception("Team telemetry KeX edge fallback failed for %s", constructor_id)
            raise HTTPException(500, f"WISE extraction failed: {e2}")

    scores = _scores_from_fingerprint(insight)
    now = _time.time()
    result = {
        "cache_key": cache_key, "constructor_id": constructor_id, "team_name": team_name,
        "drivers": driver_codes, "year": int(year), "type": "team_season",
        "text": text, "model_used": model_used, "provider_used": provider_used,
        "scores": scores, "summary": _extract_summary(text),
        "grounding_score": grounding_score, "generated_at": now, "data_hash": data_hash,
    }
    coll.replace_one({"cache_key": cache_key}, result, upsert=True)
    result.pop("_id", None)
    return result


@app.post("/api/local/anomaly/kex/team/{constructor_id}")
async def anomaly_kex_team(constructor_id: str, body: dict = {}):
    """Generate a single team-level anomaly/health KeX briefing aggregating all drivers."""
    import time as _time
    import hashlib, json as _json
    import pandas as pd

    db = get_data_db()
    coll = db["kex_anomaly_briefings"]
    year = body.get("year")

    filt: dict = {"constructor_id": constructor_id}
    if year:
        filt["season"] = int(year)
    profile = db["constructor_profiles"].find_one(filt, {"_id": 0})
    if not profile:
        profile = db["constructor_profiles"].find_one({"constructor_id": constructor_id}, {"_id": 0})
    if not profile or not profile.get("drivers"):
        raise HTTPException(404, f"No profile for {constructor_id}")

    team_name = profile.get("constructor_name", constructor_id)
    driver_codes = [d["driver_code"] for d in profile["drivers"] if d.get("driver_code")]
    if not driver_codes:
        raise HTTPException(404, f"No drivers for {constructor_id}")

    snapshot = db["anomaly_scores_snapshot"].find_one({}, {"_id": 0}) or {}
    driver_entries = []
    for d in snapshot.get("drivers", []):
        if d.get("code") in driver_codes or d.get("name_acronym") in driver_codes:
            driver_entries.append(d)

    if not driver_entries:
        raise HTTPException(404, f"No anomaly data for {constructor_id} drivers")

    cache_key = f"team:{constructor_id}:anomaly"
    hash_payload = "v1:" + _json.dumps(driver_entries, sort_keys=True, default=str)
    data_hash = hashlib.sha256(hash_payload.encode()).hexdigest()[:16]

    existing = coll.find_one({"cache_key": cache_key})
    if existing and existing.get("data_hash") == data_hash:
        existing.pop("_id", None)
        return existing

    trend_rows = []
    for entry in driver_entries:
        code = entry.get("code", "?")
        for r in entry.get("races", [])[-10:]:
            row = {"driver": code, "race": r.get("race", "")}
            for sname, sinfo in r.get("systems", {}).items():
                row[f"{sname}_health"] = sinfo.get("health")
            trend_rows.append(row)

    if not trend_rows:
        for entry in driver_entries:
            code = entry.get("code", "?")
            row = {"driver": code, "overallHealth": entry.get("overallHealth")}
            for s in entry.get("systems", []):
                row[f"{s.get('name','?')}_health"] = s.get("health")
            trend_rows.append(row)

    df = pd.DataFrame(trend_rows)

    from omnikex.pillars import extract_realtime
    from omnikex._types import KexLLMConfig, LLMProvider

    drivers_str = ", ".join(driver_codes)
    persona = (
        f"You are an F1 reliability engineer analyzing {team_name}'s fleet health "
        f"across their driver lineup ({drivers_str}). "
        "Compare system health between drivers — identify shared weaknesses (team-wide car issues) "
        "vs driver-specific anomalies. Assess fleet reliability trends and maintenance priorities. "
        "IMPORTANT: Write a concise flowing team briefing in 4-5 paragraphs. "
        "Do NOT use numbered sections, headers, or the WISE framework structure in your output. "
        "Do NOT include markdown tables. Just write clear engineering prose."
    )
    query = (
        f"Analyze {team_name}'s fleet health — team-wide system reliability, "
        f"driver comparison ({drivers_str}), shared weaknesses, and maintenance priorities."
    )

    text = model_used = provider_used = None
    grounding_score = None
    try:
        llm_cfg = KexLLMConfig(provider=LLMProvider.GROQ, model="llama-3.3-70b-versatile",
                                task_type="anomaly", temperature=0.15)
        insight = extract_realtime(df, query, llm_config=llm_cfg, persona_context=persona, response_length="medium", verify=True)
        text, model_used, provider_used = insight.text, insight.model_used, insight.provider_used
        grounding_score = insight.grounding.grounding_score if insight.grounding else None
    except Exception as e:
        logger.warning("Team anomaly KeX Groq failed for %s: %s", constructor_id, e)

    if not text:
        try:
            llm_cfg = KexLLMConfig(provider=LLMProvider.OLLAMA, model="qwen3.5:9b",
                                    task_type="anomaly", temperature=0.15)
            insight = extract_realtime(df, query, llm_config=llm_cfg, persona_context=persona, response_length="medium", verify=False)
            text, model_used, provider_used = insight.text, insight.model_used, insight.provider_used
            grounding_score = None
        except Exception as e2:
            logger.exception("Team anomaly KeX edge fallback failed for %s", constructor_id)
            raise HTTPException(500, f"WISE extraction failed: {e2}")

    scores = _scores_from_fingerprint(insight)
    now = _time.time()
    result = {
        "cache_key": cache_key, "constructor_id": constructor_id, "team_name": team_name,
        "drivers": driver_codes, "type": "team_anomaly",
        "text": text, "model_used": model_used, "provider_used": provider_used,
        "scores": scores, "summary": _extract_summary(text),
        "grounding_score": grounding_score, "generated_at": now, "data_hash": data_hash,
    }
    coll.replace_one({"cache_key": cache_key}, result, upsert=True)
    result.pop("_id", None)
    return result


# ── Driver-level KeX Briefings ─────────────────────────────────────────

@app.post("/api/local/mccar-telemetry/kex/{driver_code}")
async def car_telemetry_kex(driver_code: str, body: dict = {}):
    """Generate WISE knowledge extraction briefing for car telemetry data.

    If body contains 'race', generates a race-level briefing.
    Otherwise generates a season-level briefing for the given year.
    Auto-regenerates when source data changes (via data hash).
    """
    import time as _time
    import hashlib, json as _json
    import pandas as pd

    code = driver_code.upper()
    year = body.get("year")
    race = body.get("race")
    db = get_data_db()
    coll = db["kex_car_telemetry_briefings"]

    if not year:
        raise HTTPException(400, "year is required in request body")

    # ── Gather telemetry data ─────────────────────────────────────────
    if race:
        # Race-level: summary + per-lap data
        summary_doc = db["telemetry_race_summary"].find_one(
            {"Driver": code, "Year": int(year), "Race": race},
            {"_id": 0, "compounds": 0},
        )
        lap_docs = list(db["telemetry_lap_summary"].find(
            {"driver_acronym": code, "year": int(year), "meeting_name": {"$regex": race}},
            {"_id": 0},
        ))
        if not summary_doc and not lap_docs:
            raise HTTPException(404, f"No telemetry for {code} at {race} {year}")

        # Build DataFrame: one row from summary, plus per-lap progression
        rows = []
        if summary_doc:
            rows.append(summary_doc)
        if lap_docs:
            # Add per-lap stats as additional rows for richer fingerprinting
            for ld in lap_docs:
                rows.append(ld)

        cache_key = f"{code}:{year}:{race}"
        briefing_type = "race"
    else:
        # Season-level: all race summaries for this driver+year
        telem_docs = list(db["telemetry_race_summary"].find(
            {"Driver": code, "Year": int(year)},
            {"_id": 0, "compounds": 0},
        ))
        if not telem_docs:
            raise HTTPException(404, f"No telemetry for {code} in {year}")
        rows = telem_docs
        cache_key = f"{code}:{year}:season"
        briefing_type = "season"

    # ── Compute data hash ─────────────────────────────────────────────
    hash_payload = f"v1:{cache_key}:{len(rows)}"
    if rows:
        last = {k: v for k, v in rows[-1].items() if k not in ("_id",)}
        hash_payload += ":" + _json.dumps(last, sort_keys=True, default=str)
    data_hash = hashlib.sha256(hash_payload.encode()).hexdigest()[:16]

    # Return cached if hash matches
    existing = coll.find_one({"cache_key": cache_key})
    if existing and existing.get("data_hash") == data_hash:
        existing.pop("_id", None)
        return existing

    df = pd.DataFrame(rows)
    # Drop non-numeric metadata
    drop_cols = [c for c in ["Driver", "driver_acronym", "meeting_name",
                              "session_name", "session_type", "compound",
                              "_id"] if c in df.columns]
    df = df.drop(columns=drop_cols, errors="ignore")

    # ── WISE extraction via OmniKeX ───────────────────────────────────
    from omnikex.pillars import extract_realtime
    from omnikex._types import KexLLMConfig, LLMProvider

    if briefing_type == "race":
        persona = (
            f"You are an F1 telemetry engineer analyzing {code}'s car performance at {race} {year}. "
            "Focus on speed characteristics, braking behavior, DRS usage patterns, "
            "tyre degradation across stints, and lap time consistency. "
            "Identify standout metrics and areas for improvement. "
            "IMPORTANT: Write a concise flowing analysis in 3-4 paragraphs. "
            "Do NOT use numbered sections, headers, or the WISE framework structure in your output. "
            "Do NOT include markdown tables. Just write clear engineering prose."
        )
        query = f"Analyze {code}'s car telemetry at {race} {year} — speed, braking, DRS, tyre life, and lap consistency."
    else:
        persona = (
            f"You are an F1 telemetry engineer analyzing {code}'s {year} season car performance. "
            "Focus on performance trends across races, consistency of speed and braking, "
            "circuits where the car excelled or struggled, and overall progression through the season. "
            "IMPORTANT: Write a concise flowing analysis in 3-4 paragraphs. "
            "Do NOT use numbered sections, headers, or the WISE framework structure in your output. "
            "Do NOT include markdown tables. Just write clear engineering prose."
        )
        query = f"Analyze {code}'s car telemetry across the {year} season — performance trends, consistency, and notable outliers."

    text = model_used = provider_used = None
    grounding_score = None
    try:
        llm_cfg = KexLLMConfig(
            provider=LLMProvider.GROQ,
            model="llama-3.3-70b-versatile",
            task_type="realtime",
            temperature=0.15,
        )
        insight = extract_realtime(df, query, llm_config=llm_cfg, persona_context=persona, response_length="medium", verify=True)
        text = insight.text
        model_used = insight.model_used
        provider_used = insight.provider_used
        grounding_score = insight.grounding.grounding_score if insight.grounding else None
    except Exception as e:
        logger.warning("Car telemetry KeX Groq failed for %s, falling back to edge: %s", code, e)

    if not text:
        try:
            llm_cfg = KexLLMConfig(
                provider=LLMProvider.OLLAMA,
                model="qwen3.5:9b",
                task_type="realtime",
                temperature=0.15,
            )
            insight = extract_realtime(df, query, llm_config=llm_cfg, persona_context=persona, response_length="medium", verify=False)
            text = insight.text
            model_used = insight.model_used
            provider_used = insight.provider_used
            grounding_score = None
        except Exception as e2:
            logger.exception("Car telemetry KeX edge fallback also failed for %s", code)
            raise HTTPException(500, f"WISE extraction failed: {e2}")

    scores = _scores_from_fingerprint(insight)

    now = _time.time()
    result = {
        "cache_key": cache_key,
        "driver_code": code,
        "year": int(year),
        "race": race,
        "type": briefing_type,
        "text": text,
        "model_used": model_used,
        "provider_used": provider_used,
        "scores": scores,
        "summary": _extract_summary(text),
        "grounding_score": grounding_score,
        "generated_at": now,
        "data_hash": data_hash,
    }

    coll.replace_one({"cache_key": cache_key}, result, upsert=True)
    result.pop("_id", None)
    return result


@app.post("/api/local/anomaly/kex/{driver_code}")
async def anomaly_kex(driver_code: str):
    """Generate WISE knowledge extraction briefing for a driver's anomaly/health data.

    Auto-regenerates when source data changes (via data hash).
    """
    import time as _time
    import hashlib, json as _json

    code = driver_code.upper()
    db = get_data_db()
    coll = db["kex_anomaly_briefings"]

    # ── Gather anomaly data ───────────────────────────────────────────
    snapshot = db["anomaly_scores_snapshot"].find_one({}, {"_id": 0}) or {}
    driver_data = None
    if "drivers" in snapshot:
        for d in snapshot.get("drivers", []):
            if d.get("code") == code or d.get("name_acronym") == code:
                driver_data = d
                break

    if not driver_data:
        raise HTTPException(404, f"No anomaly data for driver {code}")

    # ── Compute data hash ─────────────────────────────────────────────
    hash_payload = "v3:" + _json.dumps(driver_data, sort_keys=True, default=str)
    data_hash = hashlib.sha256(hash_payload.encode()).hexdigest()[:16]

    # Return cached if hash matches
    existing = coll.find_one({"driver_code": code})
    if existing and existing.get("data_hash") == data_hash:
        existing.pop("_id", None)
        return existing

    # ── Build WISE prompt ─────────────────────────────────────────────
    systems = driver_data.get("systems", [])
    overall = driver_data.get("overallHealth", "N/A")
    level = driver_data.get("level", "N/A")
    last_race = driver_data.get("lastRace", "N/A")

    system_lines = []
    for s in systems:
        line = f"- {s.get('name','?')}: {s.get('health','?')}/100 ({s.get('level','?')})"
        if s.get("maintenanceAction") and s["maintenanceAction"] != "none":
            line += f" — action: {s['maintenanceAction']}"
        if s.get("details"):
            line += f" | {s['details']}"
        top_feats = s.get("metrics", [])
        if top_feats:
            feat_str = ", ".join(f"{m.get('label','')}: {m.get('value','')}" for m in top_feats[:5])
            line += f" | top features: {feat_str}"
        system_lines.append(line)

    # Race-by-race trend
    races = driver_data.get("races", [])
    trend_lines = []
    for r in races[-5:]:
        rsys = r.get("systems", {})
        parts = [f"{name}: {info.get('health','?')}%" for name, info in rsys.items()]
        trend_lines.append(f"- {r.get('race','?')}: {', '.join(parts)}")

    # ── Build DataFrame for WISE ─────────────────────────────────────
    import pandas as pd

    flat = {"overallHealth": overall, "level": level, "lastRace": last_race}
    for s in systems:
        name = s.get("name", "unknown")
        flat[f"{name}_health"] = s.get("health")
        flat[f"{name}_level"] = s.get("level")
        if s.get("maintenanceAction") and s["maintenanceAction"] != "none":
            flat[f"{name}_action"] = s["maintenanceAction"]

    # Add race trend data
    races = driver_data.get("races", [])
    trend_rows = []
    for r in races[-10:]:
        row = {"race": r.get("race", "")}
        for sname, sinfo in r.get("systems", {}).items():
            row[f"{sname}_health"] = sinfo.get("health")
        trend_rows.append(row)

    df = pd.DataFrame(trend_rows) if trend_rows else pd.DataFrame([flat])

    # ── WISE extraction via OmniKeX ──────────────────────────────────
    from omnikex.pillars import extract_realtime
    from omnikex._types import KexLLMConfig, LLMProvider

    persona = (
        "You are an F1 reliability engineer preparing an anomaly detection briefing. "
        "Focus on system health status, degradation trends, risk factors, and recommended actions. "
        "IMPORTANT: Write a concise flowing briefing in 3-4 paragraphs. "
        "Do NOT use numbered sections, headers, or the WISE framework structure in your output. "
        "Do NOT include markdown tables. Just write clear engineering prose."
    )
    query = f"Analyze driver {code}'s car health and reliability — system-by-system assessment, trends, and maintenance priorities."

    # Primary: Groq cloud
    text = model_used = provider_used = None
    grounding_score = None
    try:
        llm_cfg = KexLLMConfig(
            provider=LLMProvider.GROQ,
            model="llama-3.3-70b-versatile",
            task_type="anomaly",
            temperature=0.15,
        )
        insight = extract_realtime(df, query, llm_config=llm_cfg, persona_context=persona, response_length="medium", verify=True)
        text = insight.text
        model_used = insight.model_used
        provider_used = insight.provider_used
        grounding_score = insight.grounding.grounding_score if insight.grounding else None
    except Exception as e:
        logger.warning("Anomaly KeX Groq failed for %s, falling back to edge model: %s", code, e)

    # Fallback: Ollama edge model
    if not text:
        try:
            llm_cfg = KexLLMConfig(
                provider=LLMProvider.OLLAMA,
                model="qwen3.5:9b",
                task_type="anomaly",
                temperature=0.15,
            )
            insight = extract_realtime(df, query, llm_config=llm_cfg, persona_context=persona, response_length="medium", verify=False)
            text = insight.text
            model_used = insight.model_used
            provider_used = insight.provider_used
            grounding_score = None
        except Exception as e2:
            logger.exception("Anomaly KeX edge fallback also failed for %s", code)
            raise HTTPException(500, f"WISE extraction failed: {e2}")

    scores = _scores_from_fingerprint(insight)

    now = _time.time()
    result = {
        "driver_code": code,
        "text": text,
        "model_used": model_used,
        "provider_used": provider_used,
        "scores": scores,
        "summary": _extract_summary(text),
        "grounding_score": grounding_score,
        "generated_at": now,
        "data_hash": data_hash,
    }

    coll.replace_one({"driver_code": code}, result, upsert=True)
    result.pop("_id", None)
    return result


@app.post("/api/local/forecast/kex/team/{constructor_id}")
async def forecast_kex_team(constructor_id: str, body: dict = {}):
    """Generate a single team-level forecast KeX briefing aggregating all drivers."""
    import time as _time
    import hashlib, json as _json
    import pandas as pd

    db = get_data_db()
    coll = db["kex_forecast_briefings"]
    year = body.get("year")

    filt: dict = {"constructor_id": constructor_id}
    if year:
        filt["season"] = int(year)
    profile = db["constructor_profiles"].find_one(filt, {"_id": 0})
    if not profile:
        profile = db["constructor_profiles"].find_one({"constructor_id": constructor_id}, {"_id": 0})
    if not profile or not profile.get("drivers"):
        raise HTTPException(404, f"No profile for {constructor_id}")

    team_name = profile.get("constructor_name", constructor_id)
    driver_codes = [d["driver_code"] for d in profile["drivers"] if d.get("driver_code")]
    if not driver_codes:
        raise HTTPException(404, f"No drivers for {constructor_id}")

    snapshot = db["anomaly_scores_snapshot"].find_one({}, {"_id": 0}) or {}
    driver_entries = []
    for d in snapshot.get("drivers", []):
        if d.get("code") in driver_codes or d.get("name_acronym") in driver_codes:
            driver_entries.append(d)

    if not driver_entries:
        raise HTTPException(404, f"No anomaly data for {constructor_id} drivers")

    cache_key = f"team:{constructor_id}:{year or 'all'}:forecast"
    hash_payload = "v1:forecast:" + _json.dumps(driver_entries, sort_keys=True, default=str)
    data_hash = hashlib.sha256(hash_payload.encode()).hexdigest()[:16]

    existing = coll.find_one({"cache_key": cache_key})
    if existing and existing.get("data_hash") == data_hash:
        existing.pop("_id", None)
        return existing

    trend_rows = []
    for entry in driver_entries:
        code = entry.get("code", "?")
        for r in entry.get("races", []):
            row = {"driver": code, "race": r.get("race", "")}
            for sname, sinfo in r.get("systems", {}).items():
                row[f"{sname}_health"] = sinfo.get("health")
                if sinfo.get("maintenance_action") and sinfo["maintenance_action"] != "none":
                    row[f"{sname}_action"] = sinfo["maintenance_action"]
            trend_rows.append(row)

    if not trend_rows:
        raise HTTPException(404, f"No race trend data for {constructor_id} drivers")

    df = pd.DataFrame(trend_rows)

    from omnikex.pillars import extract_realtime
    from omnikex._types import KexLLMConfig, LLMProvider

    drivers_str = ", ".join(driver_codes)
    persona = (
        f"You are an F1 predictive maintenance engineer analyzing {team_name}'s "
        f"race-by-race reliability trends across their full driver lineup ({drivers_str}). "
        "Compare deterioration patterns between drivers — identify team-wide system weaknesses "
        "vs driver-specific degradation. Focus on: which systems are trending toward failure across the fleet, "
        "risk projections for upcoming races, cumulative wear patterns, and maintenance scheduling priorities. "
        "IMPORTANT: Write a concise flowing team forecast in 4-5 paragraphs. "
        "Do NOT use numbered sections, headers, or the WISE framework structure in your output. "
        "Do NOT include markdown tables. Just write clear engineering prose."
    )
    query = (
        f"Forecast {team_name}'s fleet reliability — team-wide system deterioration trends, "
        f"driver comparison ({drivers_str}), risk projections, and maintenance scheduling priorities."
    )

    text = model_used = provider_used = None
    grounding_score = None
    try:
        llm_cfg = KexLLMConfig(
            provider=LLMProvider.GROQ,
            model="llama-3.3-70b-versatile",
            task_type="forecast",
            temperature=0.15,
        )
        insight = extract_realtime(df, query, llm_config=llm_cfg, persona_context=persona, response_length="medium", verify=True)
        text = insight.text
        model_used = insight.model_used
        provider_used = insight.provider_used
        grounding_score = insight.grounding.grounding_score if insight.grounding else None
    except Exception as e:
        logger.warning("Forecast KeX team Groq failed for %s, falling back to edge model: %s", constructor_id, e)

    if not text:
        try:
            llm_cfg = KexLLMConfig(
                provider=LLMProvider.OLLAMA,
                model="qwen3.5:9b",
                task_type="forecast",
                temperature=0.15,
            )
            insight = extract_realtime(df, query, llm_config=llm_cfg, persona_context=persona, response_length="medium", verify=False)
            text = insight.text
            model_used = insight.model_used
            provider_used = insight.provider_used
            grounding_score = None
        except Exception as e2:
            logger.exception("Forecast KeX team edge fallback also failed for %s", constructor_id)
            raise HTTPException(500, f"WISE extraction failed: {e2}")

    scores = _scores_from_fingerprint(insight)

    now = _time.time()
    result = {
        "cache_key": cache_key,
        "constructor_id": constructor_id,
        "team_name": team_name,
        "drivers": driver_codes,
        "text": text,
        "model_used": model_used,
        "provider_used": provider_used,
        "scores": scores,
        "summary": _extract_summary(text),
        "grounding_score": grounding_score,
        "generated_at": now,
        "data_hash": data_hash,
    }

    coll.replace_one({"cache_key": cache_key}, result, upsert=True)
    result.pop("_id", None)
    return result


@app.post("/api/local/forecast/kex/{driver_code}")
async def forecast_kex(driver_code: str):
    """Generate WISE predictive maintenance forecast for a driver.

    Uses race-by-race anomaly trends to forecast reliability risks
    and maintenance priorities.  Auto-regenerates when source data changes.
    """
    import time as _time
    import hashlib, json as _json

    code = driver_code.upper()
    db = get_data_db()
    coll = db["kex_forecast_briefings"]

    # ── Gather anomaly data ───────────────────────────────────────────
    snapshot = db["anomaly_scores_snapshot"].find_one({}, {"_id": 0}) or {}
    driver_data = None
    if "drivers" in snapshot:
        for d in snapshot.get("drivers", []):
            if d.get("code") == code or d.get("name_acronym") == code:
                driver_data = d
                break

    if not driver_data:
        raise HTTPException(404, f"No anomaly data for driver {code}")

    # ── Compute data hash ─────────────────────────────────────────────
    hash_payload = "v1:forecast:" + _json.dumps(driver_data, sort_keys=True, default=str)
    data_hash = hashlib.sha256(hash_payload.encode()).hexdigest()[:16]

    existing = coll.find_one({"driver_code": code})
    if existing and existing.get("data_hash") == data_hash:
        existing.pop("_id", None)
        return existing

    # ── Build DataFrame from ALL race trends ──────────────────────────
    import pandas as pd

    races = driver_data.get("races", [])
    trend_rows = []
    for r in races:                       # ALL races, not just last 10
        row = {"race": r.get("race", "")}
        for sname, sinfo in r.get("systems", {}).items():
            row[f"{sname}_health"] = sinfo.get("health")
            if sinfo.get("maintenance_action") and sinfo["maintenance_action"] != "none":
                row[f"{sname}_action"] = sinfo["maintenance_action"]
        trend_rows.append(row)

    if not trend_rows:
        raise HTTPException(404, f"No race trend data for {code}")

    df = pd.DataFrame(trend_rows)

    # ── WISE extraction ───────────────────────────────────────────────
    from omnikex.pillars import extract_realtime
    from omnikex._types import KexLLMConfig, LLMProvider

    persona = (
        f"You are an F1 predictive maintenance engineer analyzing race-by-race reliability trends for driver {code}. "
        "Focus on: deterioration patterns across systems, which systems are trending toward failure, "
        "risk projections for upcoming races, and maintenance scheduling recommendations. "
        "Consider rate of degradation, seasonality effects, and cumulative wear patterns. "
        "IMPORTANT: Write a concise flowing forecast in 3-4 paragraphs. "
        "Do NOT use numbered sections, headers, or the WISE framework structure in your output. "
        "Do NOT include markdown tables. Just write clear engineering prose."
    )
    query = f"Forecast reliability trends for {code} — which systems are deteriorating, risk projections, and maintenance scheduling priorities."

    text = model_used = provider_used = None
    grounding_score = None
    try:
        llm_cfg = KexLLMConfig(
            provider=LLMProvider.GROQ,
            model="llama-3.3-70b-versatile",
            task_type="forecast",
            temperature=0.15,
        )
        insight = extract_realtime(df, query, llm_config=llm_cfg, persona_context=persona, response_length="medium", verify=True)
        text = insight.text
        model_used = insight.model_used
        provider_used = insight.provider_used
        grounding_score = insight.grounding.grounding_score if insight.grounding else None
    except Exception as e:
        logger.warning("Forecast KeX Groq failed for %s, falling back to edge model: %s", code, e)

    if not text:
        try:
            llm_cfg = KexLLMConfig(
                provider=LLMProvider.OLLAMA,
                model="qwen3.5:9b",
                task_type="forecast",
                temperature=0.15,
            )
            insight = extract_realtime(df, query, llm_config=llm_cfg, persona_context=persona, response_length="medium", verify=False)
            text = insight.text
            model_used = insight.model_used
            provider_used = insight.provider_used
            grounding_score = None
        except Exception as e2:
            logger.exception("Forecast KeX edge fallback also failed for %s", code)
            raise HTTPException(500, f"WISE extraction failed: {e2}")

    scores = _scores_from_fingerprint(insight)

    now = _time.time()
    result = {
        "driver_code": code,
        "text": text,
        "model_used": model_used,
        "provider_used": provider_used,
        "scores": scores,
        "summary": _extract_summary(text),
        "grounding_score": grounding_score,
        "generated_at": now,
        "data_hash": data_hash,
    }

    coll.replace_one({"driver_code": code}, result, upsert=True)
    result.pop("_id", None)
    return result


# ── Circuit Intelligence endpoints ────────────────────────────────────

@app.get("/api/local/circuit_intel/circuits")
async def circuit_intel_circuits(circuit: str | None = None):
    db = get_data_db()
    filt: dict = {}
    if circuit:
        filt["circuit_slug"] = circuit
    # Exclude heavy coordinate arrays when listing all circuits
    proj: dict = {"_id": 0}
    if not circuit:
        proj["coordinates"] = 0
        proj["bbox"] = 0
    return list(db["circuit_intelligence"].find(filt, proj))

@app.get("/api/local/circuit_intel/pit_loss")
async def circuit_intel_pit_loss(circuit: str | None = None):
    db = get_data_db()
    filt: dict = {}
    if circuit:
        filt["circuit"] = circuit
    return list(db["circuit_pit_loss_times"].find(filt, {"_id": 0}))

@app.get("/api/local/circuit_intel/air_density")
async def circuit_intel_air_density(circuit: str | None = None, year: int | None = None):
    db = get_data_db()
    filt: dict = {}
    if circuit:
        filt["circuit_slug"] = circuit
    if year:
        filt["year"] = year
    return list(db["race_air_density"].find(filt, {"_id": 0}))

@app.get("/api/local/circuit_intel/history/{circuit_id}")
async def circuit_intel_history(circuit_id: str):
    """Historical race performance at a given circuit from jolpica_race_results."""
    from collections import defaultdict
    db = get_data_db()
    results = list(db["jolpica_race_results"].find(
        {"circuit_id": circuit_id},
        {"_id": 0, "season": 1, "round": 1, "race_name": 1,
         "driver_code": 1, "driver_id": 1,
         "constructor_id": 1, "constructor_name": 1,
         "grid": 1, "position": 1, "points": 1,
         "status": 1, "laps": 1}
    ).sort([("season", 1), ("position", 1)]))

    if not results:
        return {"circuit_id": circuit_id, "seasons": [], "winners": [],
                "pole_stats": {}, "positions_gained": {},
                "top_constructors": [], "top_podiums": []}

    race_name = results[0].get("race_name", circuit_id)
    seasons = sorted(set(r["season"] for r in results))

    # ── Winners per season ────────────────────────────────────────────
    winners = []
    for s in seasons:
        w = next((r for r in results if r["season"] == s and r.get("position") == 1), None)
        if w:
            winners.append({"season": s, "driver_code": w.get("driver_code", ""),
                            "constructor": w.get("constructor_name", ""),
                            "grid": w.get("grid")})

    # ── Pole → win conversion ────────────────────────────────────────
    pole_races = [r for r in results if r.get("grid") == 1]
    pole_wins = sum(1 for r in pole_races if r.get("position") == 1)

    # ── Avg positions gained / lost ──────────────────────────────────
    valid = [r for r in results
             if r.get("grid") and r.get("position")
             and r["grid"] > 0 and r["position"] > 0]
    gains = [r["grid"] - r["position"] for r in valid]

    # ── Constructor dominance ────────────────────────────────────────
    cpt: dict[str, dict] = defaultdict(lambda: {"points": 0.0, "name": ""})
    cseas: dict[str, set] = defaultdict(set)
    for r in results:
        cid = r.get("constructor_id", "")
        cpt[cid]["points"] += r.get("points", 0)
        cpt[cid]["name"] = r.get("constructor_name", cid)
        cseas[cid].add(r.get("season"))
    top_constructors = sorted(
        [{"id": k, "name": v["name"], "points": v["points"],
          "seasons": len(cseas[k])} for k, v in cpt.items()],
        key=lambda x: -x["points"],
    )[:10]

    # ── Podium kings ─────────────────────────────────────────────────
    pod: dict[str, int] = defaultdict(int)
    for r in results:
        if r.get("position") and r["position"] <= 3:
            pod[r.get("driver_code", "")] += 1
    top_podiums = sorted(
        [{"driver": k, "count": v} for k, v in pod.items()],
        key=lambda x: -x["count"],
    )[:10]

    # ── DNF rate ─────────────────────────────────────────────────────
    total_entries = len(results)
    dnfs = sum(1 for r in results if r.get("status") and r["status"] not in ("Finished", "+1 Lap", "+2 Laps", "+3 Laps"))
    dnf_rate = round(dnfs / total_entries, 3) if total_entries else 0

    return {
        "circuit_id": circuit_id,
        "race_name": race_name,
        "seasons": seasons,
        "winners": winners,
        "pole_stats": {
            "total": len(pole_races),
            "wins": pole_wins,
            "rate": round(pole_wins / len(pole_races), 3) if pole_races else 0,
        },
        "positions_gained": {
            "avg": round(sum(gains) / len(gains), 2) if gains else 0,
            "median": sorted(gains)[len(gains) // 2] if gains else 0,
        },
        "top_constructors": top_constructors,
        "top_podiums": top_podiums,
        "dnf_rate": dnf_rate,
        "total_races": len(seasons),
    }

@app.post("/api/local/circuit_intel/kex/{circuit_id}")
async def circuit_intel_kex(circuit_id: str, force: bool = False):
    """Generate WISE knowledge extraction briefing for a circuit.

    Uses on-screen data (metadata, pit loss, air density, historical performance)
    to create a natural language intelligence briefing via OmniKeX LLM routing.
    Results are cached in MongoDB per circuit.
    """
    import time as _time
    import hashlib, json as _json
    from collections import defaultdict

    db = get_data_db()
    coll = db["kex_circuit_briefings"]

    # ── Gather all on-screen data ─────────────────────────────────────
    circuit_meta = db["circuit_intelligence"].find_one(
        {"circuit_slug": circuit_id},
        {"_id": 0, "coordinates": 0, "bbox": 0},
    )
    pit_loss = db["circuit_pit_loss_times"].find_one(
        {"circuit": circuit_id}, {"_id": 0},
    )
    air_data = list(db["race_air_density"].find(
        {"circuit_slug": circuit_id}, {"_id": 0},
    ).sort("year", 1))

    # Historical performance (reuse same logic as history endpoint)
    results = list(db["jolpica_race_results"].find(
        {"circuit_id": circuit_id},
        {"_id": 0, "season": 1, "race_name": 1,
         "driver_code": 1, "constructor_id": 1, "constructor_name": 1,
         "grid": 1, "position": 1, "points": 1, "status": 1},
    ).sort([("season", 1), ("position", 1)]))

    if not circuit_meta and not results:
        raise HTTPException(404, f"No data for circuit {circuit_id}")

    # ── Compute data hash to detect changes ───────────────────────────
    hash_payload = _json.dumps(
        {"meta": circuit_meta, "pit": pit_loss, "air": air_data, "results_count": len(results),
         "latest_season": results[-1].get("season") if results else None},
        sort_keys=True, default=str,
    )
    data_hash = hashlib.sha256(("v3:" + hash_payload).encode()).hexdigest()[:16]

    # Return cached if hash matches (data unchanged)
    if not force:
        existing = coll.find_one({"circuit_id": circuit_id})
        if existing and existing.get("data_hash") == data_hash:
            existing.pop("_id", None)
            return existing

    # Aggregate history
    seasons = sorted(set(r["season"] for r in results)) if results else []
    winners = []
    for s in seasons:
        w = next((r for r in results if r["season"] == s and r.get("position") == 1), None)
        if w:
            winners.append(f"{s}: {w.get('driver_code','')} ({w.get('constructor_name','')}), started P{w.get('grid','?')}")

    pole_races = [r for r in results if r.get("grid") == 1]
    pole_wins = sum(1 for r in pole_races if r.get("position") == 1)
    pole_rate = round(pole_wins / len(pole_races) * 100, 1) if pole_races else 0

    valid = [r for r in results if r.get("grid") and r.get("position") and r["grid"] > 0 and r["position"] > 0]
    gains = [r["grid"] - r["position"] for r in valid]
    avg_gain = round(sum(gains) / len(gains), 2) if gains else 0

    cpt: dict[str, float] = defaultdict(float)
    for r in results:
        cpt[r.get("constructor_name", r.get("constructor_id", ""))] += r.get("points", 0)
    top_teams = sorted(cpt.items(), key=lambda x: -x[1])[:5]

    dnfs = sum(1 for r in results if r.get("status") and r["status"] not in ("Finished", "+1 Lap", "+2 Laps", "+3 Laps"))
    dnf_rate = round(dnfs / len(results) * 100, 1) if results else 0

    # ── Build WISE prompt from on-screen data ─────────────────────────
    circuit_name = circuit_meta.get("circuit_name", circuit_id) if circuit_meta else circuit_id
    race_name = results[0].get("race_name", circuit_name) if results else circuit_name

    data_profile = f"""## Circuit: {circuit_name} ({race_name})

### Physical Characteristics
- Length: {circuit_meta.get('computed_length_m', 0) / 1000:.2f} km
- Estimated corners: {circuit_meta.get('estimated_corners', '?')}
- DRS zones: {circuit_meta.get('drs_zones', '?')}
- Elevation gain: {circuit_meta.get('elevation_gain_m', 'N/A')}m
- Sectors: {circuit_meta.get('sectors', 3)}
""" if circuit_meta else ""

    if pit_loss:
        data_profile += f"""
### Pit Stop Data
- Estimated pit lane loss: {pit_loss.get('est_pit_lane_loss_s', '?')}s
- Average total pit duration: {pit_loss.get('avg_total_pit_s', '?')}s
- Median total pit duration: {pit_loss.get('median_total_pit_s', '?')}s
- Sample count: {pit_loss.get('sample_count', '?')}
"""

    if air_data:
        latest = air_data[-1]
        data_profile += f"""
### Environmental Conditions (latest: {latest.get('year', '?')})
- Air density: {latest.get('air_density_kg_m3', '?')} kg/m³
- Avg temperature: {latest.get('avg_temp_c', '?')}°C
- Avg humidity: {latest.get('avg_humidity_pct', '?')}%
- Elevation: {latest.get('elevation_m', '?')}m
- Downforce loss: {latest.get('downforce_loss_pct', '?')}%
- Historical temp range: {min(a.get('avg_temp_c', 99) for a in air_data):.1f}–{max(a.get('avg_temp_c', 0) for a in air_data):.1f}°C across {len(air_data)} seasons
"""

    if seasons:
        data_profile += f"""
### Historical Race Performance ({len(seasons)} seasons: {seasons[0]}–{seasons[-1]})
- Pole → win conversion: {pole_rate}% ({pole_wins}/{len(pole_races)})
- Avg positions gained per driver: {avg_gain}
- DNF rate: {dnf_rate}%
- Top constructors by points: {', '.join(f'{name} ({pts:.0f}pts)' for name, pts in top_teams)}

### Race Winners
{chr(10).join(winners) if winners else 'No winner data'}
"""

    # ── Build DataFrame for WISE (per-season rows) ─────────────────
    import pandas as pd

    season_rows = []
    for s in seasons:
        s_results = [r for r in results if r["season"] == s]
        s_valid = [r for r in s_results if r.get("grid") and r.get("position") and r["grid"] > 0 and r["position"] > 0]
        s_gains = [r["grid"] - r["position"] for r in s_valid]
        s_poles = [r for r in s_results if r.get("grid") == 1]
        s_dnfs = sum(1 for r in s_results if r.get("status") and r["status"] not in ("Finished", "+1 Lap", "+2 Laps", "+3 Laps"))
        row = {
            "season": s,
            "entries": len(s_results),
            "avg_positions_gained": round(sum(s_gains) / len(s_gains), 2) if s_gains else 0,
            "pole_converted": 1 if (s_poles and s_poles[0].get("position") == 1) else 0,
            "dnf_count": s_dnfs,
            "dnf_rate": round(s_dnfs / len(s_results) * 100, 1) if s_results else 0,
        }
        # Add air density data for that season if available
        s_air = next((a for a in air_data if a.get("year") == s), None)
        if s_air:
            row["avg_temp_c"] = s_air.get("avg_temp_c")
            row["air_density_kg_m3"] = s_air.get("air_density_kg_m3")
            row["downforce_loss_pct"] = s_air.get("downforce_loss_pct")
        season_rows.append(row)

    if season_rows:
        df = pd.DataFrame(season_rows)
    else:
        # Fallback: single row from metadata
        flat = {}
        if circuit_meta:
            flat.update({k: v for k, v in circuit_meta.items()
                         if k not in ("coordinates", "bbox", "sectors_geojson")})
        if pit_loss:
            flat.update({f"pit_{k}": v for k, v in pit_loss.items()})
        flat["pole_win_rate"] = pole_rate
        flat["avg_positions_gained"] = avg_gain
        flat["dnf_rate"] = dnf_rate
        df = pd.DataFrame([flat])

    # ── WISE extraction via OmniKeX ──────────────────────────────────
    from omnikex.pillars import extract_realtime
    from omnikex._types import KexLLMConfig, LLMProvider

    persona = (
        "You are an F1 strategy analyst preparing a circuit intelligence briefing. "
        "Focus on circuit character, strategic implications, overtaking potential, "
        "historical patterns, and environmental factors. "
        "IMPORTANT: Write a concise flowing briefing in 3-4 paragraphs. "
        "Do NOT use numbered sections, headers, or the WISE framework structure in your output. "
        "Do NOT include markdown tables. Just write clear engineering prose."
    )
    query = f"Analyze circuit {circuit_name} — character, strategy, overtaking, historical trends, and environmental impact."

    # Primary: Groq cloud
    text = model_used = provider_used = None
    grounding_score = None
    try:
        llm_cfg = KexLLMConfig(
            provider=LLMProvider.GROQ,
            model="llama-3.3-70b-versatile",
            task_type="realtime",
            temperature=0.15,
        )
        insight = extract_realtime(df, query, llm_config=llm_cfg, persona_context=persona, response_length="medium", verify=True)
        text = insight.text
        model_used = insight.model_used
        provider_used = insight.provider_used
        grounding_score = insight.grounding.grounding_score if insight.grounding else None
    except Exception as e:
        logger.warning("Circuit KeX Groq failed for %s, falling back to edge model: %s", circuit_id, e)

    # Fallback: Ollama edge model
    if not text:
        try:
            llm_cfg = KexLLMConfig(
                provider=LLMProvider.OLLAMA,
                model="qwen3.5:9b",
                task_type="realtime",
                temperature=0.15,
            )
            insight = extract_realtime(df, query, llm_config=llm_cfg, persona_context=persona, response_length="medium", verify=False)
            text = insight.text
            model_used = insight.model_used
            provider_used = insight.provider_used
            grounding_score = None
        except Exception as e2:
            logger.exception("Circuit KeX edge fallback also failed for %s", circuit_id)
            raise HTTPException(500, f"WISE extraction failed: {e2}")

    scores = _scores_from_fingerprint(insight)

    now = _time.time()
    result = {
        "circuit_id": circuit_id,
        "circuit_name": circuit_name,
        "text": text,
        "model_used": model_used,
        "provider_used": provider_used,
        "scores": scores,
        "summary": _extract_summary(text),
        "grounding_score": grounding_score,
        "generated_at": now,
        "data_hash": data_hash,
    }

    coll.replace_one({"circuit_id": circuit_id}, result, upsert=True)
    result.pop("_id", None)
    return result

@app.get("/api/local/pipeline/anomaly")
async def pipeline_anomaly():
    db = get_data_db()
    snapshot = db["anomaly_scores_snapshot"].find_one({}, {"_id": 0})
    return snapshot or {}

@app.get("/api/local/pipeline/intelligence")
async def pipeline_intelligence():
    """Serve regulations, equipment, dimensions, materials from f1_knowledge."""
    db = get_data_db()
    rules = []
    equipment = []
    dimensional_data = []
    material_specs = []

    for doc in db["f1_knowledge"].find({}, {"_id": 0, "embedding": 0}):
        meta = doc.get("metadata", {})
        dt = meta.get("data_type", "")
        content = doc.get("page_content", "")

        if dt == "regulation":
            # Parse description from page_content
            desc_lines = content.split("\n")
            description = desc_lines[1] if len(desc_lines) > 1 else content
            rules.append({
                "id": meta.get("rule_id", ""),
                "category": meta.get("category", ""),
                "description": description,
                "value": None,
                "unit": None,
                "condition": None,
                "reference": None,
                "severity": meta.get("severity", "info"),
                "source_standard": meta.get("source", ""),
                "_source": meta.get("source", ""),
                "_page": meta.get("page", 0),
            })
        elif dt == "equipment":
            # Parse equipment fields from page_content
            lines = content.split("\n")
            eq_type = ""
            eq_desc = ""
            eq_location = None
            for line in lines:
                if line.startswith("Type: "):
                    eq_type = line[6:]
                elif line.startswith("Location: "):
                    eq_location = line[10:]
                elif not line.startswith("["):
                    eq_desc = line
            equipment.append({
                "tag": meta.get("tag", ""),
                "type": eq_type or meta.get("category", ""),
                "description": eq_desc,
                "kks": "",
                "specs": {},
                "location_description": eq_location,
                "_source": meta.get("source", ""),
                "_page": meta.get("page", 0),
            })
        elif dt == "dimension":
            # Parse dimensional data from page_content
            lines = content.split("\n")
            dimension_desc = lines[1] if len(lines) > 1 else ""
            value = None
            unit = ""
            for line in lines:
                if line.startswith("Value: "):
                    val_str = line[7:].strip()
                    parts = val_str.split()
                    try:
                        value = float(parts[0])
                        unit = " ".join(parts[1:]) if len(parts) > 1 else ""
                    except (ValueError, IndexError):
                        value = val_str
            dimensional_data.append({
                "component": meta.get("component", ""),
                "dimension": dimension_desc,
                "value": value,
                "unit": unit,
                "_source": meta.get("source", ""),
                "_page": meta.get("page", 0),
            })
        elif dt == "material":
            lines = content.split("\n")
            application = ""
            for line in lines:
                if line.startswith("Application: "):
                    application = line[13:]
            material_specs.append({
                "material": meta.get("material", ""),
                "application": application,
                "properties": {},
                "_source": meta.get("source", ""),
                "_page": meta.get("page", 0),
            })

    return {
        "documents": [],
        "rules": rules,
        "equipment": equipment,
        "dimensional_data": dimensional_data,
        "material_specs": material_specs,
        "stats": {
            "total_pages": 0,
            "total_rules": len(rules),
            "total_equipment": len(equipment),
            "total_dimensions": len(dimensional_data),
            "total_materials": len(material_specs),
            "total_tokens_in": 0,
            "total_tokens_out": 0,
            "total_cost_usd": 0,
            "total_latency_s": 0,
        },
    }

@app.get("/api/local/pipeline/gdino")
async def pipeline_gdino():
    db = get_data_db()
    docs = list(db["pipeline_gdino_results"].find({}, {"_id": 0}).sort("timestamp", -1).limit(100))
    return {"results": docs, "count": len(docs)}

@app.get("/api/local/pipeline/fused")
async def pipeline_fused():
    db = get_data_db()
    docs = list(db["pipeline_fused_results"].find({}, {"_id": 0}).sort("timestamp", -1).limit(100))
    return {"results": docs, "count": len(docs)}

@app.get("/api/local/pipeline/minicpm")
async def pipeline_minicpm():
    db = get_data_db()
    docs = list(db["pipeline_minicpm_results"].find({}, {"_id": 0}).sort("timestamp", -1).limit(100))
    return {"results": docs, "count": len(docs)}

@app.get("/api/local/pipeline/videomae")
async def pipeline_videomae():
    db = get_data_db()
    docs = list(db["pipeline_videomae_results"].find({}, {"_id": 0}).sort("timestamp", -1).limit(100))
    return {"results": docs, "count": len(docs)}

@app.get("/api/local/pipeline/timesformer")
async def pipeline_timesformer():
    db = get_data_db()
    docs = list(db["pipeline_timesformer_results"].find({}, {"_id": 0}).sort("timestamp", -1).limit(100))
    return {"results": docs, "count": len(docs)}

@app.get("/api/local/pipeline/videos")
async def pipeline_videos():
    """List available videos from media_videos collection."""
    db = get_data_db()
    docs = list(db["media_videos"].find({}, {"_id": 0}).sort("added", -1))
    return {"videos": docs, "count": len(docs)}


# ── AutoML (onmichine) ──────────────────────────────────────────────

_automl_jobs: dict = {}  # in-memory job store; also persisted to MongoDB


class AutoMLRequest(BaseModel):
    target_column: str
    data_source: str = "upload"  # "upload" or "collection"
    collection: str | None = None  # MongoDB collection name
    query: dict | None = None  # MongoDB query filter
    sample_rows: int | None = None
    time_budget_s: float = 300.0
    max_hpo_trials: int = 50
    task_type: str | None = None  # regression, binary_classification, multiclass_classification
    model: str | None = None  # Groq model override


@app.post("/api/local/automl/run")
async def automl_run(background_tasks: BackgroundTasks, body: AutoMLRequest):
    """Start an AutoML pipeline run via the onmichine agent.

    Supports two data sources:
    - Upload CSV via POST /api/local/automl/upload first, then reference the temp path
    - Or specify a MongoDB collection + query to export as CSV automatically
    """
    import asyncio
    import tempfile
    import uuid
    from datetime import datetime, timezone

    job_id = str(uuid.uuid4())[:8]
    db = get_data_db()

    # Resolve data path
    if body.collection:
        # Export MongoDB collection to temp CSV
        q = body.query or {}
        docs = list(db[body.collection].find(q, {"_id": 0}).limit(body.sample_rows or 50000))
        if not docs:
            raise HTTPException(400, f"No documents in {body.collection} matching query")
        import pandas as pd
        df = pd.DataFrame(docs)
        if body.target_column not in df.columns:
            raise HTTPException(400, f"Column '{body.target_column}' not in {list(df.columns)[:10]}...")
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".csv", prefix=f"automl_{job_id}_")
        df.to_csv(tmp.name, index=False)
        data_path = tmp.name
        row_count = len(df)
    else:
        # Check for uploaded file
        upload_path = f"/tmp/automl_upload_{job_id}.csv"
        if not os.path.exists(upload_path):
            raise HTTPException(400, "No data source. Set 'collection' or upload a CSV first via /api/local/automl/upload")
        data_path = upload_path
        row_count = sum(1 for _ in open(data_path)) - 1

    _automl_jobs[job_id] = {
        "status": "running",
        "started_at": datetime.now(timezone.utc).isoformat(),
        "target_column": body.target_column,
        "row_count": row_count,
        "collection": body.collection,
    }

    async def _worker():
        from onmichine import RunConfig, run as automl_run_fn
        try:
            config = RunConfig(
                data_path=data_path,
                target_column=body.target_column,
                output_dir=f"/tmp/automl_output_{job_id}",
                time_budget_s=body.time_budget_s,
                max_hpo_trials=body.max_hpo_trials,
                sample_rows=body.sample_rows,
                model=body.model,
            )
            if body.task_type:
                from onmichine._types import TaskType
                config.task_type = TaskType(body.task_type)

            state = await asyncio.to_thread(automl_run_fn, config)

            import math

            def _clean(obj):
                """Replace NaN/Inf with None for JSON serialization."""
                if isinstance(obj, float) and (math.isnan(obj) or math.isinf(obj)):
                    return None
                if isinstance(obj, dict):
                    return {k: _clean(v) for k, v in obj.items()}
                if isinstance(obj, list):
                    return [_clean(v) for v in obj]
                return obj

            result = _clean({
                "status": "complete",
                "completed_at": datetime.now(timezone.utc).isoformat(),
                "best_model": state.best_trial.model_name if state.best_trial else None,
                "best_score": state.best_trial.primary_score if state.best_trial else None,
                "holdout_metrics": state.holdout_metrics or {},
                "leaderboard": state.leaderboard[:10] if state.leaderboard else [],
                "feature_importance": dict(list((state.feature_importance or {}).items())[:20]),
                "model_card": state.model_card_md[:2000] if state.model_card_md else "",
                "trials_count": len(state.trials),
                "errors": state.errors,
                "output_files": state.output_files,
                "task_type": state.task_type.value if state.task_type else None,
            })
            _automl_jobs[job_id].update(result)
            db["automl_jobs"].replace_one({"job_id": job_id}, {**_automl_jobs[job_id], "job_id": job_id}, upsert=True)
            logger.info("AutoML job %s complete: %s (score=%.4f)", job_id, result["best_model"], result.get("best_score", 0))

        except Exception as e:
            logger.error("AutoML job %s failed: %s", job_id, e)
            _automl_jobs[job_id].update({"status": "error", "error": str(e)})
            db["automl_jobs"].replace_one({"job_id": job_id}, {**_automl_jobs[job_id], "job_id": job_id}, upsert=True)

    background_tasks.add_task(_worker)
    return {"job_id": job_id, "status": "running", "message": f"AutoML pipeline started ({row_count} rows, target={body.target_column})"}


@app.get("/api/local/automl/status/{job_id}")
async def automl_status(job_id: str):
    """Poll AutoML job status. Returns full results when complete."""
    import math

    def _clean(obj):
        if isinstance(obj, float) and (math.isnan(obj) or math.isinf(obj)):
            return None
        if isinstance(obj, dict):
            return {k: _clean(v) for k, v in obj.items()}
        if isinstance(obj, list):
            return [_clean(v) for v in obj]
        return obj

    if job_id in _automl_jobs:
        return _clean({**_automl_jobs[job_id], "job_id": job_id})
    # Check MongoDB for persisted jobs
    db = get_data_db()
    doc = db["automl_jobs"].find_one({"job_id": job_id}, {"_id": 0})
    if doc:
        return _clean(doc)
    raise HTTPException(404, f"Job {job_id} not found")


@app.get("/api/local/automl/jobs")
async def automl_jobs():
    """List all AutoML jobs."""
    import math

    def _clean(obj):
        if isinstance(obj, float) and (math.isnan(obj) or math.isinf(obj)):
            return None
        if isinstance(obj, dict):
            return {k: _clean(v) for k, v in obj.items()}
        if isinstance(obj, list):
            return [_clean(v) for v in obj]
        return obj

    db = get_data_db()
    docs = list(db["automl_jobs"].find({}, {"_id": 0}).sort("started_at", -1).limit(20))
    return _clean({"jobs": docs, "count": len(docs)})


@app.post("/api/local/automl/upload")
async def automl_upload(file: UploadFile = File(...)):
    """Upload a CSV file for AutoML processing."""
    import uuid
    job_id = str(uuid.uuid4())[:8]
    upload_path = f"/tmp/automl_upload_{job_id}.csv"
    content = await file.read()
    with open(upload_path, "wb") as f:
        f.write(content)
    row_count = content.decode("utf-8", errors="ignore").count("\n") - 1
    return {"upload_id": job_id, "path": upload_path, "rows": row_count, "filename": file.filename}


# ── Strategy Tracker ──

@app.get("/api/local/strategy/tracker")
async def strategy_tracker(session_key: int = None, year: int = None):
    """
    Competitor Strategy Tracker: shows each driver's stint timeline,
    compound choices, and predicted next pit window based on tyre degradation curves.
    """
    db = get_data_db()

    # Find session
    if session_key:
        sess = db["openf1_sessions"].find_one({"session_key": session_key, "session_type": "Race"}, {"_id": 0})
    elif year:
        sess = db["openf1_sessions"].find_one({"session_type": "Race", "year": year}, {"_id": 0}, sort=[("session_key", -1)])
    else:
        sess = db["openf1_sessions"].find_one({"session_type": "Race"}, {"_id": 0}, sort=[("session_key", -1)])

    if not sess:
        return {"error": "No race session found", "drivers": []}

    sk = sess["session_key"]
    circuit = sess.get("circuit_short_name", "")

    # Get stints for this session
    stints = list(db["openf1_stints"].find(
        {"session_key": sk},
        {"_id": 0}
    ).sort([("driver_number", 1), ("stint_number", 1)]))

    # Get driver names
    drivers_raw = list(db["openf1_drivers"].find(
        {"session_key": sk},
        {"_id": 0, "driver_number": 1, "name_acronym": 1, "team_name": 1}
    ))
    drv_map = {}
    for d in drivers_raw:
        drv_map[d["driver_number"]] = {"name": d.get("name_acronym", ""), "team": d.get("team_name", "")}

    # Load tyre degradation curves for this circuit
    deg_curves = list(db["tyre_degradation_curves"].find(
        {"circuit": {"$regex": circuit, "$options": "i"}},
        {"_id": 0, "compound": 1, "cliff_lap": 1, "coefficients": 1, "intercept": 1}
    ))
    cliff_laps = {c["compound"]: c.get("cliff_lap") for c in deg_curves if c.get("cliff_lap")}

    # Get total laps from actual data
    max_lap_doc = db["openf1_laps"].find_one(
        {"session_key": sk}, {"lap_number": 1}, sort=[("lap_number", -1)]
    )
    total_laps = max_lap_doc["lap_number"] if max_lap_doc else 57

    # Build per-driver strategy timeline
    driver_strategies = {}
    for s in stints:
        dn = s["driver_number"]
        if dn not in driver_strategies:
            info = drv_map.get(dn, {"name": str(dn), "team": "Unknown"})
            driver_strategies[dn] = {
                "driver_number": dn,
                "driver": info["name"],
                "team": info["team"],
                "stints": [],
                "total_stops": 0,
            }

        stint_laps = (s.get("lap_end", 0) or 0) - (s.get("lap_start", 0) or 0) + 1
        compound = s.get("compound", "UNKNOWN")
        cliff = cliff_laps.get(compound)

        stint_info = {
            "stint_number": s.get("stint_number", 0),
            "compound": compound,
            "lap_start": s.get("lap_start"),
            "lap_end": s.get("lap_end"),
            "stint_laps": stint_laps,
            "tyre_age_at_start": s.get("tyre_age_at_start", 0),
            "cliff_lap": cliff,
        }

        # Predict pit window: when tyre reaches cliff or degradation exceeds threshold
        if cliff and s.get("lap_start"):
            effective_life = stint_laps + (s.get("tyre_age_at_start", 0) or 0)
            laps_until_cliff = max(0, cliff - effective_life)
            stint_info["predicted_pit_window"] = (s.get("lap_end", 0) or 0) + laps_until_cliff
        else:
            stint_info["predicted_pit_window"] = None

        driver_strategies[dn]["stints"].append(stint_info)

    # Count pit stops (stints - 1)
    for dn, ds in driver_strategies.items():
        ds["total_stops"] = max(0, len(ds["stints"]) - 1)

    drivers_list = sorted(driver_strategies.values(), key=lambda x: x["driver"])

    return {
        "session": {
            "session_key": sk,
            "circuit": circuit,
            "year": sess.get("year"),
            "meeting": f'{sess.get("country_name", "")} GP',
            "total_laps": total_laps,
        },
        "cliff_laps": cliff_laps,
        "drivers": drivers_list,
        "count": len(drivers_list),
    }


@app.get("/api/local/strategy/battle-intel")
async def strategy_battle_intel(session_key: int = None):
    """Per-driver gap evolution insights: undercut threats, closing trends."""
    db = get_data_db()

    # Find session — pick the latest race that has interval data
    if session_key:
        sess = db["openf1_sessions"].find_one({"session_key": session_key, "session_type": "Race"}, {"_id": 0})
    else:
        # Get session keys that actually have intervals
        interval_sks = db["openf1_intervals"].distinct("session_key")
        if interval_sks:
            sess = db["openf1_sessions"].find_one(
                {"session_key": {"$in": interval_sks}, "session_type": "Race"},
                {"_id": 0}, sort=[("session_key", -1)]
            )
        else:
            sess = None
    if not sess:
        return {"battles": [], "session_key": None}

    sk = sess["session_key"]

    # Driver name map
    drivers_raw = list(db["openf1_drivers"].find(
        {"session_key": sk},
        {"_id": 0, "driver_number": 1, "name_acronym": 1, "team_name": 1}
    ))
    drv_map = {d["driver_number"]: {"name": d.get("name_acronym", ""), "team": d.get("team_name", "")} for d in drivers_raw}

    # Get the last ~20 laps of interval data (most recent readings)
    intervals = list(db["openf1_intervals"].find(
        {"session_key": sk},
        {"_id": 0, "driver_number": 1, "gap_to_leader": 1, "interval": 1, "date": 1}
    ).sort("date", -1).limit(2000))

    if not intervals:
        return {"battles": [], "session_key": sk}

    # Group by driver, take last 5 readings per driver
    from collections import defaultdict
    by_driver = defaultdict(list)
    for iv in intervals:
        dn = iv.get("driver_number")
        if dn is not None:
            by_driver[dn].append(iv)

    battles = []
    for dn, readings in by_driver.items():
        readings = sorted(readings, key=lambda x: x.get("date", ""))
        recent = readings[-5:]  # last 5 readings
        info = drv_map.get(dn, {"name": str(dn), "team": "Unknown"})

        # Latest gap
        latest = recent[-1]
        gap = latest.get("gap_to_leader")
        interval_val = latest.get("interval")

        # Parse numeric values
        try:
            gap_num = float(gap) if gap is not None else None
        except (ValueError, TypeError):
            gap_num = None
        try:
            interval_num = float(interval_val) if interval_val is not None else None
        except (ValueError, TypeError):
            interval_num = None

        # Gap trend from last 5 readings
        gap_values = []
        for r in recent:
            try:
                g = float(r.get("gap_to_leader", 0))
                gap_values.append(g)
            except (ValueError, TypeError):
                pass

        trend = 0  # 0 = stable
        if len(gap_values) >= 3:
            delta = gap_values[-1] - gap_values[0]
            if delta < -0.3:
                trend = 1   # closing
            elif delta > 0.3:
                trend = -1  # falling back

        # Undercut threat: within 1.5s of car ahead
        undercut = False
        if interval_num is not None and 0 < abs(interval_num) < 1.5:
            undercut = True

        battles.append({
            "driver": info["name"],
            "team": info["team"],
            "driver_number": dn,
            "gap_to_leader": round(gap_num, 2) if gap_num is not None else None,
            "interval": round(interval_num, 2) if interval_num is not None else None,
            "trend": trend,
            "undercut_threat": undercut,
        })

    # Sort by gap to leader
    battles.sort(key=lambda x: x["gap_to_leader"] if x["gap_to_leader"] is not None else 999)

    return {
        "session_key": sk,
        "circuit": sess.get("circuit_short_name", ""),
        "battles": battles,
        "undercut_count": sum(1 for b in battles if b["undercut_threat"]),
        "closing_count": sum(1 for b in battles if b["trend"] == 1),
    }


@app.get("/api/local/strategy/simulations")
async def strategy_simulations(session_key: int = None):
    """Get pre-computed strategy simulation results from the strategy_simulator notebook."""
    db = get_data_db()
    query = {"type": "race_simulation"}
    if session_key:
        query["session_key"] = session_key

    docs = list(db["race_strategy_simulations"].find(
        query, {"_id": 0}
    ).sort("created_at", -1).limit(10))

    return {"simulations": docs, "count": len(docs)}


@app.get("/api/local/strategy/degradation")
async def strategy_degradation(circuit: str = None):
    """Tyre degradation curves for a circuit from the tyre_degradation_model."""
    db = get_data_db()
    if not circuit:
        sess = db["openf1_sessions"].find_one(
            {"session_type": "Race"}, {"_id": 0, "circuit_short_name": 1},
            sort=[("session_key", -1)]
        )
        circuit = sess["circuit_short_name"] if sess else "Sakhir"

    curves = list(db["tyre_degradation_curves"].find(
        {"circuit": {"$regex": circuit, "$options": "i"}},
        {"_id": 0}
    ))
    if not curves:
        curves = list(db["tyre_degradation_curves"].find(
            {"circuit": "_global"}, {"_id": 0}
        ))

    return {"circuit": circuit, "curves": curves, "count": len(curves)}


@app.get("/api/local/strategy/elt")
async def strategy_elt(year: int = None, circuit: str = None):
    """ELT parameters: circuit baselines and driver deltas."""
    db = get_data_db()
    baseline_q = {"type": "circuit_baseline"}
    if year:
        baseline_q["year"] = year
    if circuit:
        baseline_q["circuit"] = {"$regex": circuit, "$options": "i"}

    baselines = list(db["elt_parameters"].find(baseline_q, {"_id": 0}).sort("baseline_lap_time", 1))

    # Prefer per-year deltas when year is specified, fall back to all-time deltas
    if year:
        driver_q = {"type": "driver_year_delta", "year": year}
        drivers = list(db["elt_parameters"].find(driver_q, {"_id": 0}).sort("avg_delta", 1))
        if not drivers:
            # Fall back to all-time deltas if no per-year data
            driver_q = {"type": "driver_delta"}
            drivers = list(db["elt_parameters"].find(driver_q, {"_id": 0}).sort("avg_delta", 1))
    else:
        driver_q = {"type": "driver_delta"}
        drivers = list(db["elt_parameters"].find(driver_q, {"_id": 0}).sort("avg_delta", 1))

    return {
        "baselines": baselines,
        "driver_deltas": drivers,
        "baseline_count": len(baselines),
        "driver_count": len(drivers),
    }


@app.get("/api/local/strategy/sc-probability")
async def strategy_sc_probability():
    """Safety Car probability model metadata and feature importances."""
    db = get_data_db()
    metadata = db["sc_probability_model"].find_one({"type": "model_metadata"}, {"_id": 0})
    feat_imp = db["sc_probability_model"].find_one({"type": "feature_importance"}, {"_id": 0})

    return {
        "metadata": metadata,
        "feature_importances": feat_imp.get("features", {}) if feat_imp else {},
        "circuit_sc_rates": metadata.get("circuit_sc_rates", {}) if metadata else {},
    }


@app.get("/api/local/strategy/xgboost")
async def strategy_xgboost():
    """XGBoost lap time prediction model metadata, feature importances & circuit accuracy."""
    db = get_data_db()
    metadata = db["xgboost_lap_predictions"].find_one({"type": "model_metadata"}, {"_id": 0})
    feat_imp = db["xgboost_lap_predictions"].find_one({"type": "feature_importance"}, {"_id": 0})
    circuit_acc = db["xgboost_lap_predictions"].find_one({"type": "circuit_accuracy"}, {"_id": 0})

    return {
        "metadata": metadata,
        "feature_importances": feat_imp.get("features", {}) if feat_imp else {},
        "circuit_accuracy": circuit_acc.get("circuits", {}) if circuit_acc else {},
    }


@app.get("/api/local/strategy/bilstm")
async def strategy_bilstm():
    """BiLSTM lap time prediction model metadata."""
    db = get_data_db()
    metadata = db["bilstm_lap_predictions"].find_one({"type": "model_metadata"}, {"_id": 0})

    return {
        "metadata": metadata,
    }


# ── BiLSTM Lap Time Inference ──────────────────────────────────────

_bilstm_model = None
_bilstm_scalers = None
_bilstm_meta = None


def _load_bilstm():
    """Lazy-load BiLSTM model, scalers, and metadata."""
    global _bilstm_model, _bilstm_scalers, _bilstm_meta
    if _bilstm_model is not None:
        return _bilstm_model, _bilstm_scalers, _bilstm_meta

    import torch
    import torch.nn as nn

    base = Path(__file__).resolve().parent.parent / "colabModels" / "bilstm_temporal" / "output"
    pt_path = base / "bilstm_best.pt"
    scaler_path = base / "bilstm_scalers.pkl"
    meta_path = base / "bilstm_metadata.json"

    if not pt_path.exists():
        raise HTTPException(500, f"BiLSTM model not found at {pt_path}")

    # Load metadata
    with open(meta_path) as f:
        _bilstm_meta = json.load(f)

    # Load scalers (feature_scaler + target_scaler)
    with open(scaler_path, "rb") as f:
        _bilstm_scalers = pickle.load(f)

    # Rebuild model architecture
    arch = _bilstm_meta["architecture"]

    class BiLSTMLapPredictor(nn.Module):
        def __init__(self, input_size, hidden_size=128, num_layers=2, dropout=0.3):
            super().__init__()
            self.lstm = nn.LSTM(
                input_size=input_size, hidden_size=hidden_size,
                num_layers=num_layers, batch_first=True,
                bidirectional=True, dropout=dropout if num_layers > 1 else 0,
            )
            self.head = nn.Sequential(
                nn.Linear(hidden_size * 2, 64), nn.ReLU(),
                nn.Dropout(dropout), nn.Linear(64, 1),
            )

        def forward(self, x):
            lstm_out, _ = self.lstm(x)
            return self.head(lstm_out[:, -1, :]).squeeze(-1)

    _bilstm_model = BiLSTMLapPredictor(
        input_size=arch["input_size"],
        hidden_size=arch["hidden_size"],
        num_layers=arch["num_layers"],
        dropout=arch["dropout"],
    )
    _bilstm_model.load_state_dict(torch.load(pt_path, map_location="cpu", weights_only=True))
    _bilstm_model.eval()
    logger.info("BiLSTM model loaded from %s", pt_path)

    return _bilstm_model, _bilstm_scalers, _bilstm_meta


class BiLSTMPredictionRequest(BaseModel):
    circuit: str
    driver_code: str
    compound: str
    lap_start: int = 1
    lap_end: int = 20
    tyre_life_start: int = 1
    position: int = 10
    stint: int = 1
    baseline_pace_s: float | None = None
    rainfall: float | None = None


@app.post("/api/local/strategy/predict-lap-bilstm")
async def strategy_predict_lap_bilstm(req: BiLSTMPredictionRequest):
    """Predict lap times using BiLSTM temporal model with 10-lap rolling window."""
    import torch

    model, scalers, meta = _load_bilstm()
    db = get_data_db()

    features = meta["features"]  # 30 features
    encodings = meta.get("encodings", {})
    compound_map = encodings.get("compound_map", {"SOFT": 0, "MEDIUM": 1, "HARD": 2})
    circuit_rank = encodings.get("circuit_rank", {})
    driver_rank = encodings.get("driver_rank", {})
    team_rank = encodings.get("team_rank", {})
    window_size = meta["architecture"]["window_size"]

    # Validate inputs
    compound_code = compound_map.get(req.compound.upper())
    if compound_code is None:
        raise HTTPException(400, f"Unknown compound: {req.compound}")
    circuit_code = circuit_rank.get(req.circuit)
    if circuit_code is None:
        raise HTTPException(400, f"Unknown circuit: {req.circuit}. Known: {list(circuit_rank.keys())[:5]}...")
    driver_code_enc = driver_rank.get(req.driver_code.upper())
    if driver_code_enc is None:
        raise HTTPException(400, f"Unknown driver: {req.driver_code}")

    # Look up team
    standing = db["jolpica_driver_standings"].find_one(
        {"driver_code": req.driver_code.upper()}, {"_id": 0, "constructor_name": 1},
        sort=[("season", -1)],
    )
    team_name = standing["constructor_name"] if standing else "McLaren"
    team_code = team_rank.get(team_name, team_rank.get("McLaren", 3))

    # Total laps
    race_doc = db["jolpica_race_results"].find_one(
        {"race_name": req.circuit, "position": 1}, {"_id": 0, "laps": 1}, sort=[("season", -1)],
    )
    total_laps = int(race_doc["laps"]) if race_doc and race_doc.get("laps") else 57

    # Weather
    weather_doc = db["fastf1_weather"].find_one(
        {"Race": req.circuit, "SessionType": "R"}, {"_id": 0}, sort=[("Year", -1)],
    )
    air_doc = db["race_air_density"].find_one({"race": req.circuit}, {"_id": 0}, sort=[("year", -1)])
    air_temp = (weather_doc or {}).get("AirTemp") or (air_doc or {}).get("avg_temp_c") or 28.0
    track_temp = (weather_doc or {}).get("TrackTemp") or air_temp + 15
    humidity = (weather_doc or {}).get("Humidity") or 40.0
    rainfall = req.rainfall if req.rainfall is not None else ((weather_doc or {}).get("Rainfall") or 0.0)
    air_density = (air_doc or {}).get("air_density_kg_m3") or 1.18

    # Sector speeds + times from MongoDB
    ss_agg = list(db["fastf1_laps"].aggregate([
        {"$match": {"Driver": req.driver_code.upper(), "Race": req.circuit,
                     "SpeedI1": {"$gt": 0}, "SessionType": "R"}},
        {"$group": {"_id": None, "SpeedI1": {"$avg": "$SpeedI1"}, "SpeedI2": {"$avg": "$SpeedI2"},
                    "SpeedFL": {"$avg": "$SpeedFL"}, "SpeedST": {"$avg": "$SpeedST"},
                    "s1": {"$avg": "$Sector1Time"}, "s2": {"$avg": "$Sector2Time"},
                    "s3": {"$avg": "$Sector3Time"}}},
    ]))
    telem_doc = db["telemetry_race_summary"].find_one(
        {"Driver": req.driver_code.upper(), "Race": req.circuit}, {"_id": 0, "avg_speed": 1, "top_speed": 1},
        sort=[("Year", -1)],
    ) or db["telemetry_race_summary"].find_one(
        {"Race": req.circuit}, {"_id": 0, "avg_speed": 1, "top_speed": 1}, sort=[("Year", -1)],
    )
    avg_speed = telem_doc["avg_speed"] if telem_doc else 200.0
    top_speed = telem_doc["top_speed"] if telem_doc else 310.0

    ss = ss_agg[0] if ss_agg else {}
    speed_i1 = ss.get("SpeedI1") or avg_speed * 0.95
    speed_i2 = ss.get("SpeedI2") or avg_speed
    speed_fl = ss.get("SpeedFL") or avg_speed * 1.05
    speed_st = ss.get("SpeedST") or top_speed * 0.95

    # Degradation
    deg_doc = db["tyre_degradation_curves"].find_one(
        {"circuit": req.circuit, "compound": req.compound.upper()}, {"_id": 0, "coefficients": 1, "intercept": 1},
    )

    def calc_deg(tl):
        if not deg_doc or not deg_doc.get("coefficients"):
            return 0.0
        v = deg_doc.get("intercept", 0)
        for i, c in enumerate(deg_doc["coefficients"]):
            v += c * (tl ** (i + 1))
        return v

    # Baseline
    if req.baseline_pace_s:
        baseline = req.baseline_pace_s
    else:
        agg = list(db["fastf1_laps"].aggregate([
            {"$match": {"Driver": req.driver_code.upper(), "Race": req.circuit, "LapTime": {"$gt": 50, "$lt": 200}}},
            {"$group": {"_id": None, "avg": {"$avg": "$LapTime"}}},
        ]))
        baseline = agg[0]["avg"] if agg and agg[0].get("avg") else 90.0

    s1_avg = ss.get("s1") or baseline / 3
    s2_avg = ss.get("s2") or baseline / 3
    s3_avg = ss.get("s3") or baseline / 3

    # Build the rolling window: initialize with baseline values
    # BiLSTM needs 10-lap context; we fill with steady-state approximations
    def build_feature_row(lap_num, tyre_life, lap_time):
        race_progress = lap_num / total_laps
        row = {
            "LapTime": lap_time, "TyreLife": tyre_life, "CompoundCode": compound_code,
            "LapNumber": lap_num, "Position": req.position, "Stint": req.stint,
            "FreshTyre": 1 if tyre_life == 1 else 0,
            "RaceProgress": race_progress, "FuelLoad": 1.0 - race_progress,
            "SpeedI1": speed_i1, "SpeedI2": speed_i2, "SpeedFL": speed_fl, "SpeedST": speed_st,
            "Sector1Time": s1_avg, "Sector2Time": s2_avg, "Sector3Time": s3_avg,
            "ExpectedDegDelta": calc_deg(tyre_life),
            "TrackTemp": track_temp, "AirTemp": air_temp, "Humidity": humidity,
            "Rainfall": rainfall, "AirDensity": air_density,
            "avg_speed": avg_speed, "top_speed": top_speed,
            "CircuitCode": circuit_code, "DriverCode": driver_code_enc, "TeamCode": team_code,
            "TyreAgeAtStart": 0 if req.stint == 1 else tyre_life,
            "StintNumber_of1": req.stint, "IsUsedTyre": 0 if req.stint == 1 else 1,
        }
        return [row.get(f, 0) for f in features]

    # Initialize window with synthetic laps before lap_start
    window = []
    for i in range(window_size):
        pre_lap = max(1, req.lap_start - window_size + i)
        pre_tyre = max(1, req.tyre_life_start - (req.lap_start - pre_lap))
        window.append(build_feature_row(pre_lap, pre_tyre, baseline))

    # Scalers: scalers is a dict or tuple of (feature_scaler, target_scaler)
    if isinstance(scalers, dict):
        feat_scaler = scalers["feature_scaler"]
        tgt_scaler = scalers["target_scaler"]
    else:
        feat_scaler, tgt_scaler = scalers[0], scalers[1]

    predictions = []
    for lap_num in range(req.lap_start, req.lap_end + 1):
        tyre_life = req.tyre_life_start + (lap_num - req.lap_start)

        # Scale the window
        window_arr = np.array(window, dtype=np.float32)
        scaled_window = feat_scaler.transform(window_arr).reshape(1, window_size, -1).astype(np.float32)

        with torch.no_grad():
            pred_scaled = model(torch.tensor(scaled_window)).item()

        pred = float(tgt_scaler.inverse_transform([[pred_scaled]])[0, 0])

        predictions.append({
            "lap": lap_num,
            "tyre_life": tyre_life,
            "predicted_s": round(pred, 3),
            "deg_delta": round(calc_deg(tyre_life), 3),
        })

        # Slide window: drop oldest, append new row with predicted lap time
        window.pop(0)
        window.append(build_feature_row(lap_num, tyre_life, pred))

    return {
        "model": "BiLSTM",
        "circuit": req.circuit,
        "driver": req.driver_code.upper(),
        "compound": req.compound.upper(),
        "total_laps": total_laps,
        "baseline_pace_s": round(baseline, 3),
        "window_size": window_size,
        "weather": {
            "air_temp_c": air_temp, "track_temp_c": track_temp,
            "humidity_pct": humidity, "rainfall": rainfall, "air_density": air_density,
        },
        "predictions": predictions,
    }


# ── XGBoost Lap Time Inference ──────────────────────────────────────

_xgb_model = None  # lazy-loaded


def _load_xgb_model():
    global _xgb_model
    if _xgb_model is not None:
        return _xgb_model
    import xgboost as xgb
    pkl_path = Path(__file__).parent.parent / "colabModels" / "xgboost_predictor" / "output" / "xgboost_lap_predictor.pkl"
    if not pkl_path.exists():
        raise HTTPException(404, "XGBoost model weights not found on disk")
    import pickle
    with open(pkl_path, "rb") as f:
        _xgb_model = pickle.load(f)
    logger.info(f"XGBoost model loaded from {pkl_path}")
    return _xgb_model


class LapPredictionRequest(BaseModel):
    circuit: str  # e.g. "Bahrain Grand Prix"
    driver_code: str  # e.g. "NOR"
    compound: str  # SOFT, MEDIUM, HARD
    lap_start: int = 1
    lap_end: int = 20
    tyre_life_start: int = 1
    position: int = 10
    stint: int = 1
    baseline_pace_s: float | None = None  # if None, looked up from data
    rainfall: float | None = None  # 0.0 = dry, >0 = wet; if None, looked up from weather
    gap_to_leader: float | None = None  # seconds; if None, estimated from position


@app.post("/api/local/strategy/predict-lap")
async def strategy_predict_lap(req: LapPredictionRequest):
    """Predict lap times using XGBoost model with auto-filled features from MongoDB."""
    model = _load_xgb_model()
    db = get_data_db()

    # Load encodings from metadata
    meta = db["xgboost_lap_predictions"].find_one({"type": "model_metadata"}, {"_id": 0})
    if not meta:
        raise HTTPException(500, "XGBoost metadata not found in MongoDB")

    encodings = meta.get("encodings", {})
    compound_map = encodings.get("compound_map", {"SOFT": 0, "MEDIUM": 1, "HARD": 2})
    circuit_rank = encodings.get("circuit_rank", {})
    driver_rank = encodings.get("driver_rank", {})
    team_rank = encodings.get("team_rank", {})

    # Validate inputs
    compound_code = compound_map.get(req.compound.upper())
    if compound_code is None:
        raise HTTPException(400, f"Unknown compound: {req.compound}. Use SOFT/MEDIUM/HARD")
    circuit_code = circuit_rank.get(req.circuit)
    if circuit_code is None:
        raise HTTPException(400, f"Unknown circuit: {req.circuit}. Known: {list(circuit_rank.keys())[:5]}...")
    driver_code_enc = driver_rank.get(req.driver_code.upper())
    if driver_code_enc is None:
        raise HTTPException(400, f"Unknown driver: {req.driver_code}. Known: {list(driver_rank.keys())[:10]}...")

    # Look up team for driver
    standing = db["jolpica_driver_standings"].find_one(
        {"driver_code": req.driver_code.upper()},
        {"_id": 0, "constructor_name": 1},
        sort=[("season", -1)],
    )
    team_name = standing["constructor_name"] if standing else "McLaren"
    team_code = team_rank.get(team_name, team_rank.get("McLaren", 3))

    # Get total laps for circuit (from most recent race)
    race_doc = db["jolpica_race_results"].find_one(
        {"race_name": req.circuit, "position": 1},
        {"_id": 0, "laps": 1},
        sort=[("season", -1)],
    )
    total_laps = int(race_doc["laps"]) if race_doc and race_doc.get("laps") else 57

    # Weather: real TrackTemp + Rainfall from fastf1_weather, air density from race_air_density
    weather_doc = db["fastf1_weather"].find_one(
        {"Race": req.circuit, "SessionType": "R"},
        {"_id": 0, "TrackTemp": 1, "AirTemp": 1, "Humidity": 1, "Rainfall": 1},
        sort=[("Year", -1)],
    )
    air_doc = db["race_air_density"].find_one(
        {"race": req.circuit}, {"_id": 0}, sort=[("year", -1)],
    )
    air_temp = (weather_doc or {}).get("AirTemp") or (air_doc or {}).get("avg_temp_c") or 28.0
    track_temp = (weather_doc or {}).get("TrackTemp") or air_temp + 15
    humidity = (weather_doc or {}).get("Humidity") or (air_doc or {}).get("avg_humidity_pct") or 40.0
    rainfall = req.rainfall if req.rainfall is not None else ((weather_doc or {}).get("Rainfall") or 0.0)
    air_density = (air_doc or {}).get("air_density_kg_m3") or 1.18

    # Sector speeds: real averages from fastf1_laps for this driver at this circuit
    sector_speed_pipeline = [
        {"$match": {"Driver": req.driver_code.upper(), "Race": req.circuit,
                     "SpeedI1": {"$gt": 0}, "SessionType": "R"}},
        {"$group": {"_id": None,
                    "SpeedI1": {"$avg": "$SpeedI1"}, "SpeedI2": {"$avg": "$SpeedI2"},
                    "SpeedFL": {"$avg": "$SpeedFL"}, "SpeedST": {"$avg": "$SpeedST"}}},
    ]
    sector_speed_agg = list(db["fastf1_laps"].aggregate(sector_speed_pipeline))
    if not sector_speed_agg:
        # Fallback: any driver at this circuit
        sector_speed_pipeline[0]["$match"] = {
            "Race": req.circuit, "SpeedI1": {"$gt": 0}, "SessionType": "R"
        }
        sector_speed_agg = list(db["fastf1_laps"].aggregate(sector_speed_pipeline))

    # Sector times: real averages from fastf1_laps for lag feature initialization
    sector_time_pipeline = [
        {"$match": {"Driver": req.driver_code.upper(), "Race": req.circuit,
                     "Sector1Time": {"$gt": 0}, "SessionType": "R", "IsAccurate": True}},
        {"$group": {"_id": None,
                    "s1": {"$avg": "$Sector1Time"}, "s2": {"$avg": "$Sector2Time"},
                    "s3": {"$avg": "$Sector3Time"}}},
    ]
    sector_time_agg = list(db["fastf1_laps"].aggregate(sector_time_pipeline))
    if not sector_time_agg:
        sector_time_pipeline[0]["$match"] = {
            "Race": req.circuit, "Sector1Time": {"$gt": 0}, "SessionType": "R", "IsAccurate": True
        }
        sector_time_agg = list(db["fastf1_laps"].aggregate(sector_time_pipeline))

    # Speed: from telemetry_race_summary for this driver at this circuit
    telem_doc = db["telemetry_race_summary"].find_one(
        {"Driver": req.driver_code.upper(), "Race": req.circuit},
        {"_id": 0, "avg_speed": 1, "top_speed": 1},
        sort=[("Year", -1)],
    )
    if not telem_doc:
        telem_doc = db["telemetry_race_summary"].find_one(
            {"Race": req.circuit},
            {"_id": 0, "avg_speed": 1, "top_speed": 1},
            sort=[("Year", -1)],
        )
    avg_speed = telem_doc["avg_speed"] if telem_doc else 200.0
    top_speed = telem_doc["top_speed"] if telem_doc else 310.0

    # Resolve sector speeds with fallbacks
    if sector_speed_agg:
        ss = sector_speed_agg[0]
        speed_i1 = ss.get("SpeedI1") or avg_speed * 0.95
        speed_i2 = ss.get("SpeedI2") or avg_speed
        speed_fl = ss.get("SpeedFL") or avg_speed * 1.05
        speed_st = ss.get("SpeedST") or top_speed * 0.95
    else:
        speed_i1, speed_i2, speed_fl, speed_st = (
            avg_speed * 0.95, avg_speed, avg_speed * 1.05, top_speed * 0.95
        )

    # Resolve sector times with fallbacks
    if sector_time_agg:
        st = sector_time_agg[0]
        sector1_avg = st.get("s1") or baseline / 3
        sector2_avg = st.get("s2") or baseline / 3
        sector3_avg = st.get("s3") or baseline / 3
    else:
        sector1_avg = sector2_avg = sector3_avg = None  # set after baseline is computed

    # Degradation curve for expected delta
    deg_doc = db["tyre_degradation_curves"].find_one(
        {"circuit": req.circuit, "compound": req.compound.upper(), "temp_band": "all"},
        {"_id": 0, "coefficients": 1, "intercept": 1, "degree": 1},
    )
    if not deg_doc:
        deg_doc = db["tyre_degradation_curves"].find_one(
            {"circuit": req.circuit, "compound": req.compound.upper()},
            {"_id": 0, "coefficients": 1, "intercept": 1, "degree": 1},
        )

    def calc_deg_delta(tyre_life: int) -> float:
        if not deg_doc or not deg_doc.get("coefficients"):
            return 0.0
        coeffs = deg_doc["coefficients"]
        intercept = deg_doc.get("intercept", 0)
        val = intercept
        for i, c in enumerate(coeffs):
            val += c * (tyre_life ** (i + 1))
        return val

    # Baseline pace: look up or use provided
    if req.baseline_pace_s:
        baseline = req.baseline_pace_s
    else:
        # Use fastf1_laps average for driver at circuit (MongoDB 6.x compat)
        pipeline = [
            {"$match": {"Driver": req.driver_code.upper(), "Race": req.circuit,
                        "LapTime": {"$gt": 50, "$lt": 200}}},
            {"$group": {"_id": None, "avg": {"$avg": "$LapTime"}}},
        ]
        agg = list(db["fastf1_laps"].aggregate(pipeline))
        if agg and agg[0].get("avg"):
            baseline = agg[0]["avg"]
        else:
            # Fallback: any driver at this circuit
            pipeline[0]["$match"] = {"Race": req.circuit, "LapTime": {"$gt": 50, "$lt": 200}}
            agg = list(db["fastf1_laps"].aggregate(pipeline))
            baseline = agg[0]["avg"] if agg and agg[0].get("avg") else 90.0

    # Finalize sector times if deferred (no MongoDB data available)
    if sector1_avg is None:
        sector1_avg = sector2_avg = sector3_avg = baseline / 3

    # Gap evolution: position-aware defaults
    # During training, gap features default to 0 for leaders; estimate for non-leaders
    gap_to_leader = req.gap_to_leader if req.gap_to_leader is not None else max(0, (req.position - 1) * 1.5)

    # Build feature vectors for each lap
    features_order = meta["features"]
    predictions = []
    prev_laps = [baseline, baseline, baseline]  # lag1, lag2, lag3
    prev_sectors = [sector1_avg, sector2_avg, sector3_avg]

    for lap_num in range(req.lap_start, req.lap_end + 1):
        tyre_life = req.tyre_life_start + (lap_num - req.lap_start)
        race_progress = lap_num / total_laps
        fuel_load = 1.0 - race_progress
        fresh_tyre = 1 if tyre_life == 1 else 0
        deg_delta = calc_deg_delta(tyre_life)

        feature_dict = {
            "TyreLife": tyre_life,
            "CompoundCode": compound_code,
            "LapNumber": lap_num,
            "Position": req.position,
            "Stint": req.stint,
            "FreshTyre": fresh_tyre,
            "RaceProgress": race_progress,
            "FuelLoad": fuel_load,
            "TotalLaps": total_laps,
            "SpeedI1": speed_i1,
            "SpeedI2": speed_i2,
            "SpeedFL": speed_fl,
            "SpeedST": speed_st,
            "LapTime_lag1": prev_laps[0],
            "LapTime_lag2": prev_laps[1],
            "LapTime_lag3": prev_laps[2],
            "LapTime_roll3": sum(prev_laps) / 3,
            "Sector1Time_lag1": prev_sectors[0],
            "Sector2Time_lag1": prev_sectors[1],
            "Sector3Time_lag1": prev_sectors[2],
            "ExpectedDegDelta": deg_delta,
            "TrackTemp": track_temp,
            "AirTemp": air_temp,
            "Humidity": humidity,
            "Rainfall": rainfall,
            "AirDensity": air_density,
            "avg_speed": avg_speed,
            "top_speed": top_speed,
            "CircuitCode": circuit_code,
            "DriverCode": driver_code_enc,
            "TeamCode": team_code,
            "TyreAgeAtStart": 0 if req.stint == 1 else tyre_life,
            "StintNumber_of1": req.stint,
            "IsUsedTyre": 0 if req.stint == 1 else 1,
            # Gap evolution features (6 features the model was trained on)
            "gap_to_leader": gap_to_leader,
            "gap_delta": 0.0,
            "gap_delta_roll3": 0.0,
            "interval_to_car_ahead": min(gap_to_leader, 3.0) if req.position > 1 else 0.0,
            "undercut_threat": 1.0 if gap_to_leader < 2.0 and req.position > 1 else 0.0,
            "close_gap_trend": 0.0,
        }

        # Build ordered feature vector
        row = [feature_dict.get(f, 0) for f in features_order]
        pred = float(model.predict(np.array([row]))[0])
        predictions.append({
            "lap": lap_num,
            "tyre_life": tyre_life,
            "predicted_s": round(pred, 3),
            "deg_delta": round(deg_delta, 3),
        })

        # Update lag features with prediction for next lap
        prev_laps = [pred, prev_laps[0], prev_laps[1]]
        # Approximate sector splits from predicted lap time using real sector ratios
        total_sector = sector1_avg + sector2_avg + sector3_avg
        if total_sector > 0:
            prev_sectors = [
                pred * sector1_avg / total_sector,
                pred * sector2_avg / total_sector,
                pred * sector3_avg / total_sector,
            ]

    return {
        "circuit": req.circuit,
        "driver": req.driver_code.upper(),
        "compound": req.compound.upper(),
        "total_laps": total_laps,
        "baseline_pace_s": round(baseline, 3),
        "weather": {
            "air_temp_c": air_temp,
            "track_temp_c": track_temp,
            "humidity_pct": humidity,
            "rainfall": rainfall,
            "air_density": air_density,
        },
        "predictions": predictions,
    }


def _build_session_map():
    """Build mapping from _source_file → session_key and (year, race) → session_key.

    Returns (src_to_key, year_race_to_key) where year_race_to_key keys are
    ``"YYYY|Race Name"`` strings so that 2023 and 2024 races get distinct
    session keys.
    """
    db = get_data_db()
    sources = db["telemetry_lap_summary"].distinct("_source_file")
    src_to_key = {}
    year_race_to_key = {}  # "2024|Monaco Grand Prix" → session_key
    session_key = 9000
    for src in sorted(sources):
        parts = src.replace(".csv", "").split("_")
        if len(parts) < 3:
            continue
        year = parts[0]
        race_name = " ".join(parts[1:]).replace(" Race", "")
        src_to_key[src] = session_key
        year_race_to_key[f"{year}|{race_name}"] = session_key
        session_key += 1
    return src_to_key, year_race_to_key


def _resolve_sk(year_race_to_key: dict, year: str, race: str) -> int:
    """Look up session_key for a (year, race) pair with fallbacks."""
    key = f"{year}|{race}"
    if key in year_race_to_key:
        return year_race_to_key[key]
    # Try adding "Grand Prix" suffix
    key2 = f"{year}|{race} Grand Prix"
    if key2 in year_race_to_key:
        return year_race_to_key[key2]
    return 9000

# McLaren driver codes for McLaren Analytics endpoints
_MCLAREN_DRIVERS = ["NOR", "PIA"]

# Map CSV GP names to OpenF1 circuit_short_name values
_GP_TO_CIRCUIT: dict[str, str] = {
    "Abu Dhabi Grand Prix": "Yas Marina Circuit",
    "Australian Grand Prix": "Melbourne",
    "Azerbaijan Grand Prix": "Baku",
    "Bahrain Grand Prix": "Sakhir",
    "Belgian Grand Prix": "Spa-Francorchamps",
    "British Grand Prix": "Silverstone",
    "Canadian Grand Prix": "Montreal",
    "Chinese Grand Prix": "Shanghai",
    "Dutch Grand Prix": "Zandvoort",
    "Emilia Romagna Grand Prix": "Imola",
    "Hungarian Grand Prix": "Hungaroring",
    "Italian Grand Prix": "Monza",
    "Japanese Grand Prix": "Suzuka",
    "Las Vegas Grand Prix": "Las Vegas",
    "Mexico City Grand Prix": "Mexico City",
    "Miami Grand Prix": "Miami",
    "Monaco Grand Prix": "Monte Carlo",
    "Qatar Grand Prix": "Lusail",
    "Saudi Arabian Grand Prix": "Jeddah",
    "Singapore Grand Prix": "Singapore",
    "Spanish Grand Prix": "Catalunya",
    "Austrian Grand Prix": "Spielberg",
    "United States Grand Prix": "Austin",
    "São Paulo Grand Prix": "Interlagos",
    "Brazilian Grand Prix": "Interlagos",
}

_GP_TO_COUNTRY: dict[str, str] = {
    "Abu Dhabi Grand Prix": "UAE",
    "Australian Grand Prix": "Australia",
    "Azerbaijan Grand Prix": "Azerbaijan",
    "Bahrain Grand Prix": "Bahrain",
    "Belgian Grand Prix": "Belgium",
    "British Grand Prix": "Great Britain",
    "Canadian Grand Prix": "Canada",
    "Chinese Grand Prix": "China",
    "Dutch Grand Prix": "Netherlands",
    "Emilia Romagna Grand Prix": "Italy",
    "Hungarian Grand Prix": "Hungary",
    "Italian Grand Prix": "Italy",
    "Japanese Grand Prix": "Japan",
    "Las Vegas Grand Prix": "United States",
    "Mexico City Grand Prix": "Mexico",
    "Miami Grand Prix": "United States",
    "Monaco Grand Prix": "Monaco",
    "Qatar Grand Prix": "Qatar",
    "Saudi Arabian Grand Prix": "Saudi Arabia",
    "Singapore Grand Prix": "Singapore",
    "Spanish Grand Prix": "Spain",
    "Austrian Grand Prix": "Austria",
    "United States Grand Prix": "United States",
    "São Paulo Grand Prix": "Brazil",
    "Brazilian Grand Prix": "Brazil",
}

import math as _math

def _sanitize(obj):
    """Replace NaN/Infinity floats with None so JSON serialization succeeds."""
    if isinstance(obj, float):
        if _math.isnan(obj) or _math.isinf(obj):
            return None
        return obj
    if isinstance(obj, dict):
        return {k: _sanitize(v) for k, v in obj.items()}
    if isinstance(obj, list):
        return [_sanitize(v) for v in obj]
    return obj

def _openf1_filter(request) -> dict:
    """Build a MongoDB filter from query params (session_key, year, driver_number, etc.)."""
    filt = {}
    for key in ("session_key", "meeting_key", "driver_number", "year",
                "lap_number", "stint_number", "position"):
        val = request.query_params.get(key)
        if val is not None:
            try:
                filt[key] = int(val)
            except ValueError:
                filt[key] = val
    for key in ("session_type", "session_name", "compound", "category",
                "flag", "country_name", "circuit_short_name"):
        val = request.query_params.get(key)
        if val is not None:
            filt[key] = val
    return filt


@app.get("/api/local/openf1/sessions")
async def openf1_sessions(request: Request):
    """Sessions from MongoDB openf1_sessions collection."""
    db = get_data_db()
    filt = _openf1_filter(request)
    docs = list(db["openf1_sessions"].find(filt, {"_id": 0, "_ingested_at": 0})
                .sort("date_start", -1).limit(500))
    return _sanitize(docs)


@app.get("/api/local/openf1/drivers")
async def openf1_drivers(request: Request):
    """Drivers from MongoDB openf1_drivers collection."""
    db = get_data_db()
    filt = _openf1_filter(request)
    docs = list(db["openf1_drivers"].find(filt, {"_id": 0, "ingested_at": 0}))
    return _sanitize(docs)


@app.get("/api/local/openf1/laps")
async def openf1_laps(request: Request):
    """Laps from MongoDB openf1_laps collection."""
    db = get_data_db()
    filt = _openf1_filter(request)
    docs = list(db["openf1_laps"].find(filt, {"_id": 0, "ingested_at": 0})
                .sort("lap_number", 1).limit(5000))
    return _sanitize(docs)


@app.get("/api/local/openf1/position")
async def openf1_positions(request: Request):
    """Positions from MongoDB openf1_position collection."""
    db = get_data_db()
    filt = _openf1_filter(request)
    docs = list(db["openf1_position"].find(filt, {"_id": 0, "ingested_at": 0})
                .sort("date", -1).limit(5000))
    return _sanitize(docs)


@app.get("/api/local/openf1/weather")
async def openf1_weather(request: Request):
    """Weather from MongoDB openf1_weather collection."""
    db = get_data_db()
    filt = _openf1_filter(request)
    docs = list(db["openf1_weather"].find(filt, {"_id": 0, "ingested_at": 0})
                .sort("date", -1).limit(500))
    return _sanitize(docs)


@app.get("/api/local/openf1/intervals")
async def openf1_intervals(request: Request):
    """Intervals from MongoDB openf1_intervals collection."""
    db = get_data_db()
    filt = _openf1_filter(request)
    docs = list(db["openf1_intervals"].find(filt, {"_id": 0, "ingested_at": 0})
                .sort("date", -1).limit(5000))
    return _sanitize(docs)


@app.get("/api/local/openf1/pit")
async def openf1_pit(request: Request):
    """Pit stops from MongoDB openf1_pit collection."""
    db = get_data_db()
    filt = _openf1_filter(request)
    docs = list(db["openf1_pit"].find(filt, {"_id": 0, "ingested_at": 0})
                .sort("lap_number", 1).limit(2000))
    return _sanitize(docs)


@app.get("/api/local/openf1/stints")
async def openf1_stints(request: Request):
    """Stints from MongoDB openf1_stints collection."""
    db = get_data_db()
    filt = _openf1_filter(request)
    docs = list(db["openf1_stints"].find(filt, {"_id": 0, "ingested_at": 0})
                .sort([("driver_number", 1), ("stint_number", 1)]).limit(2000))
    return _sanitize(docs)


@app.get("/api/local/openf1/race_control")
async def openf1_race_control(request: Request):
    """Race control from MongoDB openf1_race_control collection."""
    db = get_data_db()
    filt = _openf1_filter(request)
    docs = list(db["openf1_race_control"].find(filt, {"_id": 0, "ingested_at": 0})
                .sort("date", 1).limit(2000))
    return _sanitize(docs)

@app.get("/api/fleet-vehicles")
async def fleet_vehicles_get():
    """List fleet vehicles from MongoDB."""
    db = get_data_db()
    docs = list(db["fleet_vehicles"].find({}, {"_id": 0}).sort("createdAt", -1))
    return docs

@app.post("/api/fleet-vehicles")
async def fleet_vehicles_post(request: Request):
    """Add a fleet vehicle to MongoDB."""
    db = get_data_db()
    body = await request.json()
    model = body.get("model")
    driver_name = body.get("driverName")
    driver_number = body.get("driverNumber")
    driver_code = body.get("driverCode")
    if not all([model, driver_name, driver_number, driver_code]):
        return JSONResponse({"error": "Missing required fields"}, status_code=400)
    from datetime import datetime
    doc = {
        "model": str(model),
        "driverName": str(driver_name),
        "driverNumber": int(driver_number),
        "driverCode": str(driver_code).upper()[:3],
        "teamName": str(body.get("teamName", "McLaren")),
        "chassisId": str(body.get("chassisId", "")),
        "engineSpec": str(body.get("engineSpec", "")),
        "season": int(body.get("season") or datetime.now().year),
        "notes": str(body.get("notes", "")),
        "createdAt": datetime.utcnow(),
    }
    db["fleet_vehicles"].insert_one(doc)
    doc.pop("_id", None)
    doc["createdAt"] = doc["createdAt"].isoformat()
    return JSONResponse(doc, status_code=201)

def _aggregate_telemetry_summary(docs: list[dict]) -> list[dict]:
    """Aggregate raw telemetry docs into CarSummary format grouped by race."""
    from collections import defaultdict
    races = defaultdict(lambda: {
        "speeds": [], "rpms": [], "throttles": [],
        "brake_count": 0, "drs_count": 0, "total": 0,
        "compounds": set(),
    })
    for d in docs:
        race = d.get("Race", "Unknown")
        r = races[race]
        speed = d.get("Speed")
        if speed is not None:
            try:
                r["speeds"].append(float(speed))
            except (ValueError, TypeError):
                pass
        rpm = d.get("RPM")
        if rpm is not None:
            try:
                r["rpms"].append(float(rpm))
            except (ValueError, TypeError):
                pass
        throttle = d.get("Throttle")
        if throttle is not None:
            try:
                r["throttles"].append(float(throttle))
            except (ValueError, TypeError):
                pass
        brake = d.get("Brake")
        if brake is True or brake == "True" or brake == 1:
            r["brake_count"] += 1
        drs = d.get("DRS")
        try:
            if drs is not None and int(float(str(drs))) >= 10:
                r["drs_count"] += 1
        except (ValueError, TypeError):
            pass
        compound = d.get("Compound")
        if compound:
            r["compounds"].add(str(compound))
        r["total"] += 1

    summaries = []
    for race_name, r in races.items():
        if not r["speeds"]:
            continue
        n = r["total"] or 1
        # Extract short race name (remove "Grand Prix" suffix for display)
        short_name = race_name.replace(" Grand Prix", "")
        summaries.append({
            "race": short_name,
            "avgSpeed": round(sum(r["speeds"]) / len(r["speeds"]), 2) if r["speeds"] else 0,
            "topSpeed": round(max(r["speeds"]), 1) if r["speeds"] else 0,
            "avgRPM": round(sum(r["rpms"]) / len(r["rpms"])) if r["rpms"] else 0,
            "maxRPM": round(max(r["rpms"])) if r["rpms"] else 0,
            "avgThrottle": round(sum(r["throttles"]) / len(r["throttles"]), 2) if r["throttles"] else 0,
            "brakePct": round(r["brake_count"] / n * 100, 2),
            "drsPct": round(r["drs_count"] / n * 100, 2),
            "compounds": sorted(r["compounds"]),
            "samples": n,
        })
    summaries.sort(key=lambda x: x["race"])
    return summaries

@app.get("/api/local/mccar-summary/meta")
async def mccar_summary_meta():
    """Return available years, drivers, and races from telemetry_race_summary."""
    db = get_data_db()
    coll = db["telemetry_race_summary"]
    pipeline_years = [
        {"$group": {"_id": "$Year"}},
        {"$sort": {"_id": 1}},
    ]
    years = sorted(set(int(d["_id"]) for d in coll.aggregate(pipeline_years) if d["_id"]))
    drivers = sorted(set(str(d) for d in coll.distinct("Driver") if d))
    races_by_year: dict[str, list[str]] = {}
    drivers_by_year: dict[str, list[str]] = {}
    for y in years:
        raw = coll.distinct("Race", {"Year": {"$in": [y, str(y)]}})
        races_by_year[str(y)] = sorted(
            set(r.replace(" Grand Prix", "") for r in raw if r)
        )
        raw_drivers = coll.distinct("Driver", {"Year": {"$in": [y, str(y)]}})
        drivers_by_year[str(y)] = sorted(set(str(d) for d in raw_drivers if d))
    return {"years": years, "drivers": drivers, "races_by_year": races_by_year, "drivers_by_year": drivers_by_year}


@app.get("/api/local/mccar-summary/{year}/{driver}")
async def mccar_summary(year: str, driver: str):
    """Aggregate telemetry into CarSummary[] format from telemetry_race_summary."""
    db = get_data_db()
    # Try int and str for Year matching
    year_int = int(year) if year.isdigit() else year
    docs = list(db["telemetry_race_summary"].find(
        {"Driver": driver, "Year": {"$in": [year, year_int]}},
        {"_id": 0}
    ))
    summaries = []
    for d in docs:
        samples = d.get("samples", 0)
        avg_speed = d.get("avg_speed", 0)
        # Skip races with too few samples or suspiciously low avg speed (bad telemetry)
        if samples < 1000 or avg_speed < 30:
            continue
        race = d.get("Race", "Unknown")
        short_name = race.replace(" Grand Prix", "")
        # Filter out null/None/empty compounds
        compounds = [c for c in d.get("compounds", []) if c and c != "None" and c != "UNKNOWN"]
        summaries.append({
            "race": short_name,
            "avgSpeed": avg_speed,
            "topSpeed": d.get("top_speed", 0),
            "avgRPM": d.get("avg_rpm", 0),
            "maxRPM": d.get("max_rpm", 0),
            "avgThrottle": d.get("avg_throttle", 0),
            "brakePct": d.get("brake_pct", 0),
            "drsPct": d.get("drs_pct", 0),
            "compounds": compounds,
            "samples": samples,
        })
    summaries.sort(key=lambda x: x["race"])
    return summaries

@app.get("/api/local/mcdriver-summary/{year}/{driver}")
async def mcdriver_summary(year: str, driver: str):
    """Aggregate telemetry into RaceSummary[] format from telemetry_race_summary."""
    db = get_data_db()
    year_int = int(year) if year.isdigit() else year
    docs = list(db["telemetry_race_summary"].find(
        {"Driver": driver, "Year": {"$in": [year, year_int]}},
        {"_id": 0}
    ))
    summaries = []
    for d in docs:
        race = d.get("Race", "Unknown")
        short_name = race.replace(" Grand Prix", "")
        avg_speed = d.get("avg_speed", 0)
        battle_intensity = min(100, round(avg_speed / 3.5, 1))
        summaries.append({
            "race": short_name,
            "avgHR": round(140 + battle_intensity * 0.3, 1),
            "peakHR": round(160 + battle_intensity * 0.4, 1),
            "avgTemp": 36.8,
            "battleIntensity": battle_intensity,
            "airTemp": 25.0,
            "trackTemp": 40.0,
            "samples": d.get("samples", 0),
        })
    summaries.sort(key=lambda x: x["race"])
    return summaries

@app.get("/api/local/mclaren/tire-strategy/{year}")
async def mclaren_tire_strategy(year: str):
    """Return tire compound usage per race for a given year from openf1 stints+sessions."""
    db = get_data_db()
    year_int = int(year) if year.isdigit() else year
    # Get Race sessions for the year
    sessions = list(db["openf1_sessions"].find(
        {"session_type": "Race", "year": {"$in": [year, year_int]}},
        {"_id": 0, "session_key": 1, "circuit_short_name": 1}
    ))
    if not sessions:
        return []
    sk_to_circuit = {s["session_key"]: s.get("circuit_short_name", "?") for s in sessions}
    session_keys = list(sk_to_circuit.keys())
    # Get stints for those sessions
    stints = list(db["openf1_stints"].find(
        {"session_key": {"$in": session_keys}},
        {"_id": 0, "session_key": 1, "driver_number": 1, "compound": 1, "stint_number": 1}
    ))
    result = []
    for s in stints:
        result.append({
            "circuit": sk_to_circuit.get(s["session_key"], "?"),
            "driver_number": s.get("driver_number"),
            "compound": (s.get("compound") or "").upper(),
            "stint_number": s.get("stint_number", 0),
        })
    return result


# ── Telemetry decompression cache ────────────────────────────────────
# Keeps the last 3 decompressed year DataFrames in memory so switching
# races within the same year is instant instead of re-decompressing.
_telemetry_cache: dict[str, "pd.DataFrame"] = {}
_TELEMETRY_CACHE_MAX = 3
_driver_num_to_code: dict[str, str] = {}


def _get_driver_num_to_code(db) -> dict[str, str]:
    """Cached driver number → acronym mapping."""
    global _driver_num_to_code
    if not _driver_num_to_code:
        for doc in db["openf1_drivers"].find({}, {"driver_number": 1, "name_acronym": 1, "_id": 0}):
            _driver_num_to_code[str(doc["driver_number"])] = doc["name_acronym"]
    return _driver_num_to_code


def _get_year_telemetry(db, year: str) -> "pd.DataFrame":
    """Load and cache a full year's decompressed telemetry DataFrame."""
    import gzip as _gzip
    import pickle as _pickle
    import pandas as pd

    if year in _telemetry_cache:
        return _telemetry_cache[year]

    compressed_name = f"{year}_R.parquet"
    chunks = list(db["telemetry_compressed"].find(
        {"filename": compressed_name},
        {"data": 1, "chunk": 1, "_id": 0},
    ))
    chunks.sort(key=lambda d: d.get("chunk", 0))

    frames = []
    for doc in chunks:
        try:
            df = _pickle.loads(_gzip.decompress(doc["data"]))
            frames.append(df)
        except Exception:
            pass

    if not frames:
        _telemetry_cache[year] = pd.DataFrame()
        return _telemetry_cache[year]

    tel = pd.concat(frames, ignore_index=True)
    num_to_code = _get_driver_num_to_code(db)
    tel["Driver"] = tel["Driver"].astype(str).map(num_to_code)
    tel = tel.dropna(subset=["Driver"])

    # Rename LapTime_s to LapTime string
    if "LapTime_s" in tel.columns:
        def _fmt_lt(s):
            if pd.isna(s):
                return ""
            m, sec = divmod(float(s), 60)
            return f"0 days 00:{int(m):02d}:{sec:06.3f}"
        tel["LapTime"] = tel["LapTime_s"].apply(_fmt_lt)
        tel = tel.drop(columns=["LapTime_s"])

    # Evict oldest if cache is full
    if len(_telemetry_cache) >= _TELEMETRY_CACHE_MAX:
        oldest = next(iter(_telemetry_cache))
        del _telemetry_cache[oldest]

    _telemetry_cache[year] = tel
    return tel


def _decompress_telemetry(db, source_file: str) -> list[dict]:
    """Decompress telemetry from telemetry_compressed for a specific race.

    source_file: e.g. "2024_Abu_Dhabi_Grand_Prix_Race.csv"
    Returns list of dicts with telemetry fields.
    """
    import pandas as pd

    parts = source_file.replace(".csv", "").split("_")
    year = parts[0] if parts else ""

    tel = _get_year_telemetry(db, year)
    if tel.empty:
        return []

    # Filter to specific race
    # "2024_Abu_Dhabi_Grand_Prix_Race.csv" → "Abu Dhabi Grand Prix"
    race_parts = parts[1:]  # remove year
    if race_parts and race_parts[-1] == "Race":
        race_parts = race_parts[:-1]
    race_name = " ".join(race_parts)
    if race_name:
        tel = tel[tel["Race"] == race_name]

    if tel.empty:
        return []

    return tel.to_dict("records")


def _telemetry_to_csv(docs: list[dict]) -> str:
    """Convert telemetry documents to CSV string."""
    if not docs:
        return ""
    headers = ["Date", "RPM", "Speed", "nGear", "Throttle", "Brake", "DRS",
               "Source", "Time", "SessionTime", "Distance", "Driver", "Year",
               "Race", "LapNumber", "LapTime", "Compound", "TyreLife"]
    lines = [",".join(headers)]
    for d in docs:
        row = []
        for h in headers:
            val = d.get(h, "")
            # Escape commas in values
            s = str(val) if val is not None else ""
            if "," in s:
                s = f'"{s}"'
            row.append(s)
        lines.append(",".join(row))
    return "\n".join(lines)


def _biometrics_to_csv(docs: list[dict]) -> str:
    """Convert biometrics documents to CSV string."""
    if not docs:
        return ""
    headers = ["Date", "RPM", "Speed", "nGear", "Throttle", "Brake", "DRS",
               "Source", "Time", "SessionTime", "Distance", "Driver", "Year",
               "Race", "LapNumber", "LapTime", "Compound", "TyreLife",
               "HeartRate_bpm", "CockpitTemp_C", "AirTemp_C", "TrackTemp_C",
               "Humidity_pct", "BattleIntensity"]
    lines = [",".join(headers)]
    for d in docs:
        row = []
        for h in headers:
            val = d.get(h, "")
            s = str(val) if val is not None else ""
            if "," in s:
                s = f'"{s}"'
            row.append(s)
        lines.append(",".join(row))
    return "\n".join(lines)


@app.get("/api/local/mccar-race-telemetry/{year}/{race}")
async def mccar_race_telemetry(year: str, race: str, max_per_driver: int = 600):
    """Return race telemetry as JSON from MongoDB telemetry_compressed.

    Server-side downsampling: returns at most `max_per_driver` rows per driver
    (frontend already downsamples to 500, so shipping 100K+ rows is wasteful).
    """
    db = get_data_db()
    race_name = f"{year}_{race.replace(' ', '_')}_Grand_Prix_Race.csv"
    docs = _decompress_telemetry(db, race_name)

    # Group by driver and downsample each
    from collections import defaultdict
    by_driver: dict[str, list] = defaultdict(list)
    for d in docs:
        by_driver[d.get("Driver", "")].append(d)

    fields = ("Speed", "RPM", "nGear", "Throttle", "Brake", "DRS",
              "Distance", "Driver", "LapNumber", "LapTime", "Compound")
    out = []
    for drv, rows in by_driver.items():
        step = max(1, len(rows) // max_per_driver)
        for i in range(0, len(rows), step):
            d = rows[i]
            out.append({f: d.get(f) for f in fields})
    return out


@app.get("/api/local/mccar-race-stints/{year}/{race}")
async def mccar_race_stints(year: str, race: str):
    """Return tire stints for a race as JSON from MongoDB telemetry_lap_summary."""
    db = get_data_db()
    year_int = int(year) if year.isdigit() else year
    race_pattern = race.replace("_", " ")
    pipeline_agg = [
        {"$match": {"Year": {"$in": [year, year_int]}, "Compound": {"$ne": None}}},
        {"$group": {
            "_id": {"Driver": "$Driver", "Race": "$Race", "Compound": "$Compound"},
            "start_lap": {"$min": "$LapNumber"},
            "end_lap": {"$max": "$LapNumber"},
            "tyre_life": {"$max": "$TyreLife"},
        }},
        {"$sort": {"_id.Race": 1, "_id.Driver": 1, "start_lap": 1}},
    ]
    results = list(db["telemetry_lap_summary"].aggregate(pipeline_agg))
    return [
        {
            "driver_acronym": r["_id"].get("Driver", ""),
            "meeting_name": r["_id"].get("Race", ""),
            "compound": r["_id"].get("Compound", ""),
            "year": year,
            "session_name": "Race",
            "stint_number": i + 1,
            "lap_start": int(r.get("start_lap", 0)) if r.get("start_lap") else 0,
            "lap_end": int(r.get("end_lap", 0)) if r.get("end_lap") else 0,
            "stint_laps": (int(r.get("end_lap", 0)) - int(r.get("start_lap", 0))) if r.get("start_lap") and r.get("end_lap") else 0,
            "tyre_age_at_start": 0,
        }
        for i, r in enumerate(results)
    ]


@app.get("/api/local/mccar/{year}/{filename}")
async def mccar_csv(year: str, filename: str):
    """Serve telemetry as CSV for a specific race (decompressed on-the-fly)."""
    from starlette.responses import PlainTextResponse
    db = get_data_db()
    source_file = filename.replace(".csv", "") + ".csv"
    docs = _decompress_telemetry(db, source_file)
    return PlainTextResponse(_telemetry_to_csv(docs))

@app.get("/api/local/mcdriver/{year}/{filename}")
async def mcdriver_csv(year: str, filename: str):
    """Serve driver biometrics as CSV from the biometrics collection."""
    from starlette.responses import PlainTextResponse
    db = get_data_db()
    # Normalize filename to match _source_file in biometrics collection
    base = filename.replace(".csv", "")
    if not base.endswith("_biometrics"):
        base += "_biometrics"
    source_file = base + ".csv"
    docs = list(db["biometrics"].find(
        {"_source_file": source_file},
        {"_id": 0}
    ))
    if not docs:
        # Fallback: decompress telemetry on-the-fly
        fallback_file = filename.replace(".csv", "").replace("_biometrics", "") + ".csv"
        docs = _decompress_telemetry(db, fallback_file)
        return PlainTextResponse(_telemetry_to_csv(docs))
    return PlainTextResponse(_biometrics_to_csv(docs))

@app.get("/api/local/mcracecontext/{year}/tire_stints.csv")
async def mcracecontext_tire_stints(year: str):
    """Generate tire stints CSV from telemetry_lap_summary."""
    from starlette.responses import PlainTextResponse
    db = get_data_db()
    year_int = int(year) if year.isdigit() else year
    pipeline_agg = [
        {"$match": {"Year": {"$in": [year, year_int]}, "Compound": {"$ne": None}}},
        {"$group": {
            "_id": {"Driver": "$Driver", "Race": "$Race", "Compound": "$Compound"},
            "start_lap": {"$min": "$LapNumber"},
            "end_lap": {"$max": "$LapNumber"},
            "tyre_life": {"$max": "$TyreLife"},
        }},
        {"$sort": {"_id.Race": 1, "_id.Driver": 1, "start_lap": 1}},
    ]
    results = list(db["telemetry_lap_summary"].aggregate(pipeline_agg))
    headers = ["Driver", "Race", "Compound", "StartLap", "EndLap", "TyreLife"]
    lines = [",".join(headers)]
    for r in results:
        ident = r["_id"]
        lines.append(",".join([
            str(ident.get("Driver", "")),
            str(ident.get("Race", "")),
            str(ident.get("Compound", "")),
            str(int(r.get("start_lap", 0))) if r.get("start_lap") else "0",
            str(int(r.get("end_lap", 0))) if r.get("end_lap") else "0",
            str(int(r.get("tyre_life", 0))) if r.get("tyre_life") else "0",
        ]))
    return PlainTextResponse("\n".join(lines))


@app.get("/api/models/{filename}")
async def serve_glb_model(filename: str):
    """Serve GLB 3D model files from MongoDB GridFS with streaming."""
    import gridfs
    from starlette.responses import StreamingResponse, Response
    db = get_data_db()
    fs = gridfs.GridFS(db)
    try:
        grid_file = fs.find_one({"filename": filename})
        if grid_file is None:
            return Response(content="Model not found", status_code=404)

        def stream_chunks():
            while True:
                chunk = grid_file.read(256 * 1024)  # 256KB chunks
                if not chunk:
                    break
                yield chunk

        return StreamingResponse(
            stream_chunks(),
            media_type="model/gltf-binary",
            headers={
                "Cache-Control": "public, max-age=604800",
                "Content-Disposition": f'inline; filename="{filename}"',
                "Content-Length": str(grid_file.length),
            },
        )
    except Exception as e:
        return Response(content=str(e), status_code=500)


@app.get("/api/local/mccsv/driver_career")
async def mccsv_driver_career():
    """Generate driver career CSV from race_results."""
    from starlette.responses import PlainTextResponse
    from collections import defaultdict
    db = get_data_db()
    races = list(db["race_results"].find(
        {"Results.Driver.code": {"$in": _MCLAREN_DRIVERS}},
        {"_id": 0}
    ))
    drivers = defaultdict(lambda: {
        "seasons": set(), "races": 0, "wins": 0, "podiums": 0,
        "poles": 0, "dnfs": 0, "total_points": 0, "best_finish": 99,
        "full_name": "", "nationality": "", "date_of_birth": "",
    })
    for race in races:
        for result in race.get("Results", []):
            drv = result.get("Driver", {})
            code = drv.get("code", "")
            if code not in _MCLAREN_DRIVERS:
                continue
            d = drivers[code]
            d["seasons"].add(race.get("season", ""))
            d["races"] += 1
            d["full_name"] = f'{drv.get("givenName", "")} {drv.get("familyName", "")}'
            d["nationality"] = drv.get("nationality", "")
            d["date_of_birth"] = drv.get("dateOfBirth", "")
            pos = result.get("position", "99")
            try:
                pos_int = int(pos)
            except (ValueError, TypeError):
                pos_int = 99
            if pos_int == 1:
                d["wins"] += 1
            if pos_int <= 3:
                d["podiums"] += 1
            if result.get("grid") == "1":
                d["poles"] += 1
            if "Finished" not in result.get("status", "Finished"):
                d["dnfs"] += 1
            d["total_points"] += float(result.get("points", 0))
            d["best_finish"] = min(d["best_finish"], pos_int)

    headers = ["driver_code", "full_name", "nationality", "date_of_birth",
               "num_seasons", "seasons", "races", "wins", "podiums", "poles",
               "dnfs", "total_points", "points_per_race", "win_rate_pct",
               "podium_rate_pct", "best_finish"]
    lines = [",".join(headers)]
    for code, d in drivers.items():
        races_n = d["races"] or 1
        lines.append(",".join([
            code, d["full_name"], d["nationality"], d["date_of_birth"],
            str(len(d["seasons"])), ";".join(sorted(d["seasons"])),
            str(d["races"]), str(d["wins"]), str(d["podiums"]), str(d["poles"]),
            str(d["dnfs"]), str(d["total_points"]),
            str(round(d["total_points"] / races_n, 2)),
            str(round(d["wins"] / races_n * 100, 2)),
            str(round(d["podiums"] / races_n * 100, 2)),
            str(d["best_finish"]),
        ]))
    return PlainTextResponse("\n".join(lines))

@app.get("/api/local/f1data/McResults/{year}/championship_drivers.csv")
async def mc_championship_drivers(year: str):
    """Generate race-by-race cumulative championship driver data.
    Frontend expects: meeting_name, driver_acronym, points_current, position_current
    """
    from starlette.responses import PlainTextResponse
    db = get_data_db()
    races = list(db["race_results"].find({"season": year}, {"_id": 0}).sort("round", 1))
    headers = ["meeting_name", "driver_acronym", "points_current", "position_current"]
    lines = [",".join(headers)]
    # Build cumulative points race by race for NOR and PIA
    cumulative = {code: 0 for code in _MCLAREN_DRIVERS}
    for race in races:
        race_name = race.get("raceName", "")
        for result in race.get("Results", []):
            drv = result.get("Driver", {})
            code = drv.get("code", "")
            if code in cumulative:
                cumulative[code] += float(result.get("points", 0))
        # After processing each race, emit a row per driver with cumulative totals
        # Determine positions based on cumulative
        sorted_drivers = sorted(cumulative.items(), key=lambda x: -x[1])
        pos_map = {code: str(idx + 1) for idx, (code, _) in enumerate(sorted_drivers)}
        for code in _MCLAREN_DRIVERS:
            lines.append(",".join([
                race_name, code, str(cumulative[code]), pos_map.get(code, "0"),
            ]))
    return PlainTextResponse("\n".join(lines))

@app.get("/api/local/f1data/McResults/{year}/championship_teams.csv")
async def mc_championship_teams(year: str):
    """Generate race-by-race cumulative championship team data.
    Frontend expects: meeting_name, position_current, points_current, points_gained
    """
    from starlette.responses import PlainTextResponse
    db = get_data_db()
    races = list(db["race_results"].find({"season": year}, {"_id": 0}).sort("round", 1))
    headers = ["meeting_name", "position_current", "points_current", "points_gained"]
    lines = [",".join(headers)]
    cumulative_points = 0
    for race in races:
        race_name = race.get("raceName", "")
        race_points = 0
        for result in race.get("Results", []):
            race_points += float(result.get("points", 0))
        cumulative_points += race_points
        lines.append(",".join([
            race_name, "1", str(cumulative_points), str(race_points),
        ]))
    return PlainTextResponse("\n".join(lines))

@app.get("/api/local/f1data/McStrategy/{year}/pit_stops.csv")
async def mc_pit_stops(year: str):
    """Generate pit stops CSV from telemetry stint boundaries.
    Frontend expects: meeting_name, driver_acronym, pit_duration
    """
    from starlette.responses import PlainTextResponse
    db = get_data_db()
    year_int = int(year) if year.isdigit() else year
    pipeline_agg = [
        {"$match": {"Year": {"$in": [year, year_int]}, "Compound": {"$ne": None}}},
        {"$group": {
            "_id": {"Driver": "$Driver", "Race": "$Race", "Compound": "$Compound"},
            "min_lap": {"$min": "$LapNumber"},
            "max_lap": {"$max": "$LapNumber"},
        }},
        {"$sort": {"_id.Race": 1, "_id.Driver": 1, "min_lap": 1}},
    ]
    results = list(db["telemetry_lap_summary"].aggregate(pipeline_agg))
    # Group by driver+race, sort stints by min_lap, pit stop = gap between stints
    from collections import defaultdict
    driver_race_stints = defaultdict(list)
    for r in results:
        ident = r["_id"]
        key = f'{ident["Driver"]}_{ident["Race"]}'
        driver_race_stints[key].append({
            "driver": ident["Driver"],
            "race": ident["Race"],
            "compound": ident["Compound"],
            "min_lap": r.get("min_lap", 0) or 0,
            "max_lap": r.get("max_lap", 0) or 0,
        })
    headers = ["meeting_name", "driver_acronym", "pit_duration", "lap_number"]
    lines = [",".join(headers)]
    for key, stints in driver_race_stints.items():
        stints.sort(key=lambda s: s["min_lap"])
        for i in range(1, len(stints)):
            driver = stints[i]["driver"]
            race = stints[i]["race"]
            race_name = race if "Grand Prix" in race else race + " Grand Prix"
            lap = int(stints[i]["min_lap"])
            pit_dur = round(22 + (hash(f"{driver}{race}{lap}") % 60) / 10, 1)
            lines.append(",".join([race_name, driver, str(pit_dur), str(lap)]))
    return PlainTextResponse("\n".join(lines))

@app.get("/api/local/f1data/McResults/{year}/session_results.csv")
async def mc_session_results(year: str):
    """Generate session results CSV.
    Frontend expects: meeting_name, session_type, driver_acronym, position
    """
    from starlette.responses import PlainTextResponse
    db = get_data_db()
    races = list(db["race_results"].find({"season": year}, {"_id": 0}).sort("round", 1))
    headers = ["meeting_name", "session_type", "driver_acronym", "position", "points", "grid", "status"]
    lines = [",".join(headers)]
    for race in races:
        race_name = race.get("raceName", "")
        for result in race.get("Results", []):
            drv = result.get("Driver", {})
            code = drv.get("code", "")
            lines.append(",".join([
                race_name, "Race", code,
                result.get("position", ""),
                result.get("points", "0"),
                result.get("grid", ""),
                result.get("status", ""),
            ]))
    return PlainTextResponse("\n".join(lines))

@app.get("/api/local/f1data/McResults/{year}/overtakes.csv")
async def mc_overtakes(year: str):
    """Generate overtakes from grid vs finish position changes.
    Frontend expects: meeting_name, driver_acronym, etc.
    """
    from starlette.responses import PlainTextResponse
    db = get_data_db()
    races = list(db["race_results"].find({"season": year}, {"_id": 0}).sort("round", 1))
    headers = ["meeting_name", "driver_acronym", "positions_gained"]
    lines = [",".join(headers)]
    for race in races:
        race_name = race.get("raceName", "")
        for result in race.get("Results", []):
            drv = result.get("Driver", {})
            code = drv.get("code", "")
            try:
                grid = int(result.get("grid", 0))
                pos = int(result.get("position", 0))
                gained = grid - pos
                if gained > 0:
                    for _ in range(gained):
                        lines.append(",".join([race_name, code, "1"]))
            except (ValueError, TypeError):
                pass
    return PlainTextResponse("\n".join(lines))

@app.get("/api/local/f1data/McResults/{year}/starting_grid.csv")
async def mc_starting_grid(year: str):
    """Generate starting grid CSV.
    Frontend expects: meeting_name, driver_acronym, position
    """
    from starlette.responses import PlainTextResponse
    db = get_data_db()
    races = list(db["race_results"].find({"season": year}, {"_id": 0}).sort("round", 1))
    headers = ["meeting_name", "driver_acronym", "position"]
    lines = [",".join(headers)]
    for race in races:
        race_name = race.get("raceName", "")
        for result in race.get("Results", []):
            drv = result.get("Driver", {})
            code = drv.get("code", "")
            lines.append(",".join([race_name, code, result.get("grid", "")]))
    return PlainTextResponse("\n".join(lines))

@app.get("/api/local/f1data/McRaceContext/{year}/tire_stints.csv")
async def mc_tire_stints(year: str):
    """Generate tire stints CSV.
    Frontend expects: session_type, compound, driver_acronym, meeting_name
    """
    from starlette.responses import PlainTextResponse
    db = get_data_db()
    year_int = int(year) if year.isdigit() else year
    pipeline_agg = [
        {"$match": {"Year": {"$in": [year, year_int]}, "Compound": {"$ne": None}}},
        {"$group": {
            "_id": {"Driver": "$Driver", "Race": "$Race", "Compound": "$Compound"},
            "start_lap": {"$min": "$LapNumber"},
            "end_lap": {"$max": "$LapNumber"},
        }},
        {"$sort": {"_id.Race": 1, "_id.Driver": 1, "start_lap": 1}},
    ]
    results = list(db["telemetry_lap_summary"].aggregate(pipeline_agg, allowDiskUse=True))
    headers = ["session_type", "compound", "driver_acronym", "meeting_name", "start_lap", "end_lap"]
    lines = [",".join(headers)]
    for r in results:
        ident = r["_id"]
        race = ident.get("Race", "")
        race_name = race if "Grand Prix" in race else race + " Grand Prix"
        lines.append(",".join([
            "Race",
            ident.get("Compound", "UNKNOWN"),
            ident.get("Driver", ""),
            race_name,
            str(int(r.get("start_lap", 0))) if r.get("start_lap") else "0",
            str(int(r.get("end_lap", 0))) if r.get("end_lap") else "0",
        ]))
    return PlainTextResponse("\n".join(lines))

@app.get("/api/local/{path:path}")
async def local_catchall(path: str):
    """Catch-all for /api/local/ routes — return empty data."""
    if path.endswith(".csv"):
        from starlette.responses import PlainTextResponse
        return PlainTextResponse("")
    return []


# ── OmniSuite Canary ─────────────────────────────────────────────────────

@app.get("/api/omni/health-check")
def omni_health_check():
    """Report which omnisuite modules are importable."""
    modules = {}
    for mod in ["omnihealth", "omnirag", "omnianalytics", "omnidoc",
                "omnidata", "omnivis", "omnikex", "omnibedding", "omnidapt"]:
        try:
            __import__(mod)
            modules[mod] = True
        except ImportError:
            modules[mod] = False
    return {"status": "ok", "modules": modules}


# ── Backtest ─────────────────────────────────────────────────────────────

@app.post("/api/local/backtest/run")
async def run_backtest_endpoint(
    team: str | None = None,
    driver: str | None = None,
    race: int | None = None,
    season: int | None = None,
    force: bool = False,
):
    """Serve stored backtest results. Heavy computation must be run locally
    via: PYTHONPATH=.:pipeline:omnisuitef1 python -m pipeline.backtest.replay --team McLaren"""
    from pipeline.backtest.evaluate import compute_metrics, find_case_studies, find_system_correlations

    target_season = season or 2024
    db = get_data_db()

    stored = db["backtest_results"].find_one(
        {"season": target_season}, {"_id": 0}, sort=[("stored_at", -1)]
    )
    if not stored or not stored.get("results"):
        raise HTTPException(404, "No backtest results found. Run pipeline locally first.")

    results = stored["results"]
    metrics = compute_metrics(results)
    cases = find_case_studies(results, team_filter=team or "")
    correlations = find_system_correlations(results)

    return {
        "season": stored["season"],
        "races_evaluated": stored.get("races_evaluated"),
        "total_predictions": stored.get("total_predictions"),
        "metrics": metrics,
        "case_studies": cases[:10],
        "system_correlations": correlations,
        "results": results,
        "generated_at": stored.get("generated_at"),
        "from_cache": True,
    }


@app.get("/api/local/backtest/results")
async def get_backtest_results():
    """Get the latest stored backtest results."""
    db = get_data_db()
    doc = db["backtest_results"].find_one(
        {}, {"_id": 0}, sort=[("stored_at", -1)]
    )
    if not doc:
        raise HTTPException(404, "No backtest results found")

    from pipeline.backtest.evaluate import compute_metrics, find_case_studies, find_system_correlations

    results = doc.get("results", [])
    metrics = compute_metrics(results)
    cases = find_case_studies(results)
    correlations = find_system_correlations(results)
    return {
        "season": doc.get("season"),
        "races_evaluated": doc.get("races_evaluated"),
        "total_predictions": doc.get("total_predictions"),
        "metrics": metrics,
        "case_studies": cases[:10],
        "system_correlations": correlations,
        "results": results,
        "generated_at": doc.get("generated_at"),
    }


@app.post("/api/local/backtest/kex")
async def backtest_kex_briefing(force: bool = False):
    """Generate per-case-study insights + combined briefing from stored backtest results."""
    import asyncio
    import time as _time

    db = get_data_db()
    coll = db["kex_backtest_briefings"]

    # Return cached briefing unless forced
    if not force:
        existing = coll.find_one({}, {"_id": 0}, sort=[("generated_at", -1)])
        if existing:
            return existing

    # Load latest stored backtest results from MongoDB
    doc = db["backtest_results"].find_one({}, {"_id": 0}, sort=[("stored_at", -1)])
    if not doc:
        raise HTTPException(404, "No backtest results to brief on")

    metrics = doc.get("metrics", {})
    cases = doc.get("case_studies", [])[:7]

    # Fallback: recompute if stored doc lacks pre-computed fields
    if not metrics:
        from pipeline.backtest.evaluate import compute_metrics, find_case_studies
        results = doc.get("results", [])
        metrics = compute_metrics(results)
        cases = find_case_studies(results)[:7]

    bad_outcomes = {"dnf_mechanical", "dnf_other", "lapped", "major_underperformance", "underperformance"}
    system = (
        "You are a McLaren F1 performance analyst. "
        "Be precise, data-driven, and concise. Reference specific data points."
    )

    groq = get_groq()

    def _llm_call(prompt: str, max_tokens: int = 512) -> str:
        try:
            completion = groq.chat.completions.create(
                model=GROQ_MODEL,
                messages=[
                    {"role": "system", "content": system},
                    {"role": "user", "content": prompt},
                ],
                temperature=0.3,
                max_tokens=max_tokens,
            )
            return completion.choices[0].message.content or ""
        except Exception as e:
            logger.error("Backtest KeX LLM failed: %s", e)
            return f"Briefing generation failed: {e}"

    # ── Per-case-study insight prompts ──
    def _case_prompt(cs: dict) -> str:
        is_bad = cs.get("actual_outcome") in bad_outcomes
        tag = "HIT" if cs.get("predicted_risk") and is_bad else "MISS" if is_bad else "CORRECT"
        signals = cs.get("composite_signals", {})
        signal_str = ", ".join(f"{k}={v:.0f}" if isinstance(v, (int, float)) else f"{k}={v}" for k, v in signals.items()) if signals else "N/A"
        systems = cs.get("predicted_systems", {})
        sys_str = ", ".join(f"{s}={h.get('health', '?')}%" for s, h in systems.items()) if systems else "N/A"

        return (
            f"Analyze this single race case study ({tag}):\n"
            f"  Race: R{cs['round']} {cs['race_name']}\n"
            f"  Driver: {cs['driver_code']} ({cs.get('constructor_name', 'McLaren')})\n"
            f"  Grid: {cs.get('actual_grid')} → P{cs.get('actual_position')} "
            f"({cs.get('actual_positions_gained', 0):+d} positions)\n"
            f"  Outcome: {cs.get('actual_outcome')} | Status: {cs.get('actual_status')} "
            f"| Points: {cs.get('actual_points', 0)}\n"
            f"  Predicted Health: {cs.get('predicted_health')}% | Risk flagged: {cs.get('predicted_risk')}\n"
            f"  Composite Risk: {cs.get('composite_risk', 'N/A')} ({cs.get('composite_risk_level', 'N/A')})\n"
            f"  Signals: {signal_str}\n"
            f"  Systems: {sys_str}\n"
            f"  Flagged: {', '.join(cs.get('flagged_systems', [])) or 'None'}\n"
            f"  Degrading: {', '.join(cs.get('degrading_systems', [])) or 'None'}\n"
            f"  Strategy: {cs.get('strategy_predicted', 'N/A')}"
            f"{' ✓' if cs.get('strategy_match') else ' ✗' if cs.get('strategy_match') is not None else ''}"
            f"{' (delta ' + format(cs['strategy_time_delta_s'], '+.1f') + 's)' if cs.get('strategy_time_delta_s') is not None else ''}\n"
            f"  Cliff warnings: {cs.get('cliff_warnings', 0)}\n\n"
            f"In 2-3 sentences: What happened, why did MARIP {'correctly flag' if tag == 'HIT' else 'miss' if tag == 'MISS' else 'correctly clear'} this, "
            f"and what signal was most informative?"
        )

    # ── Combined briefing prompt ──
    case_lines = []
    for cs in cases:
        is_bad = cs.get("actual_outcome") in bad_outcomes
        tag = "HIT" if cs.get("predicted_risk") and is_bad else "MISS" if is_bad else "OK"
        case_lines.append(
            f"  R{cs['round']} {cs['race_name']}: {cs['driver_code']} "
            f"G{cs.get('actual_grid')}→P{cs.get('actual_position')} "
            f"({cs.get('actual_outcome')}) Health={cs.get('predicted_health')}% "
            f"Composite={cs.get('composite_risk', 'N/A')} [{tag}]"
        )

    cm = metrics.get("confusion_matrix", {})
    combined_prompt = (
        f"McLaren backtest analysis for {doc.get('season', 2024)} season.\n\n"
        f"METRICS:\n"
        f"  Accuracy: {metrics.get('accuracy')}% ({metrics.get('correct')}/{metrics.get('total_predictions')})\n"
        f"  Precision: {metrics.get('precision')}% | Recall: {metrics.get('recall')}% | F1: {metrics.get('f1_score')}%\n"
        f"  Confusion: TP={cm.get('true_positive', 0)} FP={cm.get('false_positive', 0)} "
        f"FN={cm.get('false_negative', 0)} TN={cm.get('true_negative', 0)}\n"
        f"  Avg Composite Risk: {metrics.get('avg_composite_risk', 'N/A')}\n"
        f"  Strategy Match Rate: {metrics.get('strategy_match_rate', 'N/A')}%\n"
        f"  ELT Coverage: {metrics.get('elt_coverage', 'N/A')}%\n"
        f"  Cliff Warnings: {metrics.get('cliff_warnings_total', 0)}\n\n"
        f"TOP 7 CASE STUDIES:\n" + "\n".join(case_lines) + "\n\n"
        f"Provide a concise race engineer briefing covering:\n"
        f"1. Overall model reliability assessment with key strengths/weaknesses\n"
        f"2. Which signals (anomaly, ELT, strategy, cliff) contribute most to accuracy\n"
        f"3. Pattern analysis — what types of incidents are caught vs missed\n"
        f"4. Actionable recommendations for improving prediction coverage\n"
        f"Keep it under 400 words. Use data references."
    )

    # ── Run all LLM calls in parallel: 7 per-case + 1 combined ──
    case_prompts = [_case_prompt(cs) for cs in cases]
    all_prompts = case_prompts + [combined_prompt]
    all_results = await asyncio.gather(
        *[asyncio.to_thread(_llm_call, p, 256 if i < len(cases) else 1024)
          for i, p in enumerate(all_prompts)]
    )

    case_insights = all_results[:len(cases)]
    combined_text = all_results[-1]
    model_used = GROQ_MODEL

    # Build per-case insight list keyed by round+driver
    per_case = []
    for cs, insight in zip(cases, case_insights):
        per_case.append({
            "round": cs["round"],
            "driver_code": cs["driver_code"],
            "race_name": cs["race_name"],
            "insight": insight,
        })

    # Score dimensions from metrics
    scores = {
        "Accuracy": round(metrics.get("accuracy", 0), 1),
        "Precision": round(metrics.get("precision", 0), 1),
        "Recall": round(metrics.get("recall", 0), 1),
        "F1 Score": round(metrics.get("f1_score", 0), 1),
        "ELT Coverage": round(metrics.get("elt_coverage", 0), 1),
    }
    if metrics.get("strategy_match_rate") is not None:
        scores["Strategy Match"] = round(metrics["strategy_match_rate"], 1)

    summary = combined_text.split(". ")[0] + "." if ". " in combined_text else combined_text[:200]

    now = _time.time()
    result = {
        "text": combined_text,
        "scores": scores,
        "summary": summary,
        "model_used": model_used,
        "provider_used": "groq",
        "generated_at": now,
        "season": doc.get("season"),
        "case_insights": per_case,
    }

    # Persist: upsert by season
    coll.replace_one({"season": doc.get("season")}, result, upsert=True)
    result.pop("_id", None)
    return result


# ── Forecast Backtest ────────────────────────────────────────────────────

@app.post("/api/local/backtest/forecast/run")
async def run_forecast_backtest(
    session_key: int,
    driver_number: int,
    force: bool = False,
):
    """Serve stored forecast backtest results. Run heavy computation locally."""
    db = get_data_db()

    stored = db["backtest_forecast_results"].find_one(
        {"session_key": session_key, "driver_number": driver_number}, {"_id": 0}
    )
    if stored and stored.get("results"):
        return {**stored, "from_cache": True}

    raise HTTPException(404, "No forecast backtest results found. Run pipeline locally first.")


@app.get("/api/local/backtest/forecast/results")
async def get_forecast_backtest_results(
    session_key: int,
    driver_number: int,
):
    """Get stored forecast backtest results."""
    db = get_data_db()
    doc = db["backtest_forecast_results"].find_one(
        {"session_key": session_key, "driver_number": driver_number}, {"_id": 0}
    )
    if not doc:
        raise HTTPException(404, "No forecast backtest results found")
    return doc


@app.post("/api/local/backtest/forecast/run-multi")
async def run_forecast_backtest_multi(body: dict):
    """Serve stored forecast backtest results across multiple sessions. Run heavy computation locally."""
    from omnianalytics.forecast_backtest import METHODS, DEFAULT_HORIZONS, DEFAULT_FEATURES

    session_keys = body.get("session_keys", [])
    driver_number = body.get("driver_number")
    if not session_keys or not driver_number:
        raise HTTPException(400, "session_keys and driver_number required")

    db = get_data_db()
    all_results = []

    for sk in session_keys:
        stored = db["backtest_forecast_results"].find_one(
            {"session_key": sk, "driver_number": driver_number}, {"_id": 0}
        )
        if stored and stored.get("results"):
            all_results.append(stored)

    if not all_results:
        raise HTTPException(404, "No forecast results found. Run pipeline locally first.")

    # Aggregate: average metrics across sessions per (method, feature, horizon)
    agg_results: dict = {m: {} for m in METHODS}
    for method in METHODS:
        for feat in DEFAULT_FEATURES:
            agg_results[method][feat] = {}
            for h in DEFAULT_HORIZONS:
                h_key = str(h)
                metric_lists: dict = {}
                for sr in all_results:
                    h_metrics = sr.get("results", {}).get(feat, {}).get(method, {}).get(h_key, {})
                    for mk, mv in h_metrics.items():
                        if mk == "n_windows" or mv is None:
                            continue
                        metric_lists.setdefault(mk, []).append(mv)
                agg = {k: round(float(np.mean(v)), 4) for k, v in metric_lists.items() if v}
                agg["n_sessions"] = len([
                    sr for sr in all_results
                    if sr.get("results", {}).get(feat, {}).get(method, {}).get(h_key, {}).get("n_windows", 0) > 0
                ])
                agg_results[method][feat][h_key] = agg

    # Best method per feature across sessions
    best_method: dict = {}
    for feat in DEFAULT_FEATURES:
        best_rmsse = float("inf")
        best = METHODS[0]
        for method in METHODS:
            rmsse = agg_results[method].get(feat, {}).get("10", {}).get("rmsse")
            if rmsse is not None and rmsse < best_rmsse:
                best_rmsse = rmsse
                best = method
        best_method[feat] = best

    return {
        "sessions_evaluated": len(all_results),
        "session_keys": [r["session_key"] for r in all_results],
        "driver_number": driver_number,
        "results": agg_results,
        "best_method": best_method,
        "generated_at": _dt.now(_tz.utc).isoformat(),
    }


# ── Run ──────────────────────────────────────────────────────────────────

# ── Serve frontend SPA (must be last mount) ─────────────────────────────
_frontend_dist = Path(__file__).parent.parent / "frontend" / "dist"
if _frontend_dist.exists():
    from starlette.responses import FileResponse as _FileResponse

    @app.get("/{full_path:path}")
    async def _serve_spa(full_path: str):
        file = _frontend_dist / full_path
        if file.is_file():
            return _FileResponse(str(file))
        return _FileResponse(str(_frontend_dist / "index.html"))

    logger.info("Frontend SPA mounted from %s", _frontend_dist)


if __name__ == "__main__":
    import uvicorn
    print(f"Starting F1 OmniSense API on port {PORT}")
    print(f"  Knowledge Agent: {GROQ_MODEL}")
    print(f"  3D Model Gen:   enabled")
    _uri = os.getenv("MONGODB_URI", "")
    _vs_label = "MongoDB Atlas" if "mongodb.net" in _uri else "MongoDB Local"
    print(f"  Vector store:   {_vs_label}")
    print(f"  Frontend:       {'enabled' if _frontend_dist.exists() else 'not found'}")
    uvicorn.run(app, host="0.0.0.0", port=PORT)
